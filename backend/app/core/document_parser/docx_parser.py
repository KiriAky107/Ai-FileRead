"""
Word 文档 (.docx) 解析器
"""
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

from .base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Word 文档解析器"""

    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
        self.parser_name = "docx_parser"

    def parse(
        self,
        file_path: str,
        **kwargs
    ) -> ParseResult:
        """
        解析 Word 文档

        Args:
            file_path: 文件路径
            **kwargs: 其他参数

        Returns:
            ParseResult: 解析结果
        """
        path = Path(file_path)

        # 检查文件是否存在
        if not path.exists():
            return ParseResult(
                success=False,
                error=f"文件不存在: {file_path}"
            )

        # 尝试使用 python-docx 解析，失败则使用备用方法
        try:
            return self._parse_with_docx(path)
        except Exception as e:
            logger.warning(f"python-docx 解析失败，使用备用方法: {e}")
            try:
                return self._parse_fallback(path)
            except Exception as fallback_error:
                logger.error(f"备用解析方法也失败: {fallback_error}")
                return ParseResult(
                    success=False,
                    error=f"解析 Word 文档失败: {str(e)}"
                )

    def _parse_with_docx(self, path: Path) -> ParseResult:
        """使用 python-docx 解析文档"""
        # 检查文件扩展名
        if path.suffix.lower() not in self.supported_extensions:
            return ParseResult(
                success=False,
                error=f"不支持的文件类型: {path.suffix}"
            )

        # 读取 Word 文档
        doc = Document(path)

        # 提取文本内容
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    "text": para.text,
                    "style": str(para.style.name) if para.style else "Normal"
                })

        # 提取段落纯文本（用于 AI 解析）
        paragraphs_text = [p["text"] for p in paragraphs if p["text"].strip()]

        # 提取表格内容
        tables_data = []
        for i, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_data)

            if table_rows:
                # 第一行作为表头，其余行作为数据
                headers = table_rows[0] if table_rows else []
                data_rows = table_rows[1:] if len(table_rows) > 1 else []
                tables_data.append({
                    "table_index": i,
                    "headers": headers,  # 添加 headers 字段
                    "rows": data_rows,  # 数据行（不含表头）
                    "row_count": len(data_rows),
                    "column_count": len(headers) if headers else 0
                })

        # 提取图片/嵌入式对象信息
        images_info = self._extract_images_info(doc, path)

        # 合并所有文本（包括图片描述）
        full_text_parts = []
        full_text_parts.append("【文档正文】")
        full_text_parts.extend(paragraphs_text)

        if tables_data:
            full_text_parts.append("\n【文档表格】")
            for idx, table in enumerate(tables_data):
                full_text_parts.append(f"--- 表格 {idx + 1} ---")
                for row in table["rows"]:
                    full_text_parts.append(" | ".join(str(cell) for cell in row))

        if images_info.get("image_count", 0) > 0:
            full_text_parts.append(f"\n【文档图片】文档包含 {images_info['image_count']} 张图片/图表")

        full_text = "\n".join(full_text_parts)

        # 构建元数据
        metadata = {
            "filename": path.name,
            "extension": path.suffix.lower(),
            "paragraph_count": len(paragraphs),
            "table_count": len(tables_data),
            "image_count": images_info.get("image_count", 0)
        }

        return ParseResult(
            success=True,
            data={
                "content": full_text,
                "paragraphs": paragraphs,
                "paragraphs_with_style": paragraphs,
                "tables": tables_data,
                "images": images_info
            },
            metadata=metadata
        )

    def _parse_fallback(self, path: Path) -> ParseResult:
        """备用解析方法：直接解析 docx 的 XML 结构"""
        import zipfile
        from xml.etree import ElementTree as ET

        try:
            with zipfile.ZipFile(path, 'r') as zf:
                # 读取 document.xml
                if 'word/document.xml' not in zf.namelist():
                    return ParseResult(success=False, error="无效的 docx 文件格式")

                xml_content = zf.read('word/document.xml')
                root = ET.fromstring(xml_content)

                # 命名空间
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }

                paragraphs = []
                tables = []
                current_table = []

                for elem in root.iter():
                    if elem.tag.endswith('}p'):  # 段落
                        text_parts = []
                        for t in elem.iter():
                            if t.tag.endswith('}t') and t.text:
                                text_parts.append(t.text)
                        text = ''.join(text_parts).strip()
                        if text:
                            paragraphs.append({'text': text, 'style': 'Normal'})
                    elif elem.tag.endswith('}tr'):  # 表格行
                        row_data = []
                        for tc in elem.iter():
                            if tc.tag.endswith('}tc'):  # 单元格
                                cell_text = []
                                for t in tc.iter():
                                    if t.tag.endswith('}t') and t.text:
                                        cell_text.append(t.text)
                                row_data.append(''.join(cell_text).strip())
                        if row_data:
                            current_table.append(row_data)
                    else:
                        # 表格结束，保存
                        if current_table:
                            tables.append({
                                'table_index': len(tables),
                                'rows': current_table,
                                'row_count': len(current_table),
                                'column_count': len(current_table[0]) if current_table else 0
                            })
                            current_table = []

                # 保存最后一张表格
                if current_table:
                    tables.append({
                        'table_index': len(tables),
                        'rows': current_table,
                        'row_count': len(current_table),
                        'column_count': len(current_table[0]) if current_table else 0
                    })

                # 构建文本
                paragraphs_text = [p["text"] for p in paragraphs]
                full_text_parts = ["【文档正文】"] + paragraphs_text

                if tables:
                    full_text_parts.append("\n【文档表格】")
                    for idx, table in enumerate(tables):
                        full_text_parts.append(f"--- 表格 {idx + 1} ---")
                        for row in table["rows"]:
                            full_text_parts.append(" | ".join(str(cell) for cell in row))

                full_text = "\n".join(full_text_parts)

                return ParseResult(
                    success=True,
                    data={
                        "content": full_text,
                        "paragraphs": paragraphs,
                        "paragraphs_with_style": paragraphs,
                        "tables": tables,
                        "images": {"image_count": 0, "descriptions": []}
                    },
                    metadata={
                        "filename": path.name,
                        "extension": path.suffix.lower(),
                        "paragraph_count": len(paragraphs),
                        "table_count": len(tables),
                        "image_count": 0,
                        "parse_method": "fallback_xml"
                    }
                )

        except zipfile.BadZipFile:
            return ParseResult(success=False, error="无效的 ZIP/文档文件")
        except Exception as e:
            return ParseResult(success=False, error=f"备用解析失败: {str(e)}")

    def extract_images_as_base64(self, file_path: str) -> List[Dict[str, str]]:
        """
        提取 Word 文档中的所有图片，返回 base64 编码列表

        Args:
            file_path: Word 文件路径

        Returns:
            图片列表，每项包含 base64 编码和图片类型
        """
        import zipfile
        import base64
        from io import BytesIO

        images = []

        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                # 查找 word/media 目录下的图片文件
                for filename in zf.namelist():
                    if filename.startswith('word/media/'):
                        # 获取图片类型
                        ext = filename.split('.')[-1].lower()
                        mime_types = {
                            'png': 'image/png',
                            'jpg': 'image/jpeg',
                            'jpeg': 'image/jpeg',
                            'gif': 'image/gif',
                            'bmp': 'image/bmp'
                        }
                        mime_type = mime_types.get(ext, 'image/png')

                        try:
                            # 读取图片数据并转为 base64
                            image_data = zf.read(filename)
                            base64_data = base64.b64encode(image_data).decode('utf-8')

                            images.append({
                                "filename": filename,
                                "mime_type": mime_type,
                                "base64": base64_data,
                                "size": len(image_data)
                            })
                            logger.info(f"提取图片: {filename}, 大小: {len(image_data)} bytes")
                        except Exception as e:
                            logger.warning(f"提取图片失败 {filename}: {str(e)}")

        except Exception as e:
            logger.error(f"打开 Word 文档提取图片失败: {str(e)}")

        logger.info(f"共提取 {len(images)} 张图片")
        return images

    def extract_text_from_images(self, file_path: str, lang: str = 'chi_sim+eng') -> Dict[str, Any]:
        """
        对 Word 文档中的图片进行 OCR 文字识别

        Args:
            file_path: Word 文件路径
            lang: Tesseract 语言代码，默认简体中文+英文 (chi_sim+eng)

        Returns:
            包含识别结果的字典
        """
        import zipfile
        from io import BytesIO
        from PIL import Image

        try:
            import pytesseract
        except ImportError:
            logger.warning("pytesseract 未安装，OCR 功能不可用")
            return {
                "success": False,
                "error": "pytesseract 未安装，请运行: pip install pytesseract",
                "image_count": 0,
                "extracted_text": []
            }

        results = {
            "success": True,
            "image_count": 0,
            "extracted_text": [],
            "total_chars": 0
        }

        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                # 查找 word/media 目录下的图片文件
                media_files = [f for f in zf.namelist() if f.startswith('word/media/')]

                for idx, filename in enumerate(media_files):
                    ext = filename.split('.')[-1].lower()
                    if ext not in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                        continue

                    try:
                        # 读取图片数据
                        image_data = zf.read(filename)
                        image = Image.open(BytesIO(image_data))

                        # 使用 Tesseract OCR 提取文字
                        text = pytesseract.image_to_string(image, lang=lang)
                        text = text.strip()

                        if text:
                            results["extracted_text"].append({
                                "image_index": idx,
                                "filename": filename,
                                "text": text,
                                "char_count": len(text)
                            })
                            results["total_chars"] += len(text)

                        logger.info(f"图片 {filename} OCR 识别完成，提取 {len(text)} 字符")

                    except Exception as e:
                        logger.warning(f"图片 {filename} OCR 识别失败: {str(e)}")

                results["image_count"] = len(results["extracted_text"])

        except zipfile.BadZipFile:
            results["success"] = False
            results["error"] = "无效的 Word 文档文件"
        except Exception as e:
            results["success"] = False
            results["error"] = f"OCR 处理失败: {str(e)}"

        return results

    def extract_key_sentences(self, text: str, max_sentences: int = 10) -> List[str]:
        """
        从文本中提取关键句子

        Args:
            text: 文本内容
            max_sentences: 最大句子数

        Returns:
            关键句子列表
        """
        # 简单实现：按句号分割，取前N个句子
        sentences = [s.strip() for s in text.split("。") if s.strip()]
        return sentences[:max_sentences]

    def extract_structured_fields(self, text: str) -> Dict[str, Any]:
        """
        尝试提取结构化字段

        针对合同、简历等有固定格式的文档

        Args:
            text: 文本内容

        Returns:
            提取的字段字典
        """
        fields = {}

        # 常见字段模式
        patterns = {
            "姓名": r"姓名[：:]\s*(\S+)",
            "电话": r"电话[：:]\s*(\d{11}|\d{3}-\d{8})",
            "邮箱": r"邮箱[：:]\s*(\S+@\S+)",
            "地址": r"地址[：:]\s*(.+?)(?:\n|$)",
            "金额": r"金额[：:]\s*(\d+(?:\.\d+)?)",
            "日期": r"日期[：:]\s*(\d{4}[年/-]\d{1,2}[月/-]\d{1,2})",
        }

        import re
        for field_name, pattern in patterns.items():
            match = re.search(pattern, text)
            if match:
                fields[field_name] = match.group(1)

        return fields

    def parse_tables_for_template(
        self,
        file_path: str
    ) -> Dict[str, Any]:
        """
        解析 Word 文档中的表格，提取模板字段

        专门用于比赛场景：解析表格模板，识别需要填写的字段

        Args:
            file_path: Word 文件路径

        Returns:
            包含表格字段信息的字典
        """
        from docx import Document
        from docx.table import Table
        from docx.oxml.ns import qn

        doc = Document(file_path)

        template_info = {
            "tables": [],
            "fields": [],
            "field_count": 0
        }

        for table_idx, table in enumerate(doc.tables):
            table_info = {
                "table_index": table_idx,
                "rows": [],
                "headers": [],
                "data_rows": [],
                "field_hints": {}  # 字段名称 -> 提示词/描述
            }

            # 提取表头（第一行）
            if table.rows:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                table_info["headers"] = header_cells

                # 提取数据行
                for row_idx, row in enumerate(table.rows[1:], 1):
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_info["data_rows"].append(row_data)
                    table_info["rows"].append({
                        "row_index": row_idx,
                        "cells": row_data
                    })

                # 尝试从第二列/第三列提取提示词
                # 比赛模板通常格式为：字段名 | 提示词 | 填写值
                if len(table.rows[0].cells) >= 2:
                    for row_idx, row in enumerate(table.rows[1:], 1):
                        cells = [cell.text.strip() for cell in row.cells]
                        if len(cells) >= 2 and cells[0]:
                            # 第一列是字段名
                            field_name = cells[0]
                            # 第二列可能是提示词或描述
                            hint = cells[1] if len(cells) > 1 else ""
                            table_info["field_hints"][field_name] = hint

                            template_info["fields"].append({
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "field_name": field_name,
                                "hint": hint,
                                "expected_value": cells[2] if len(cells) > 2 else ""
                            })

            template_info["tables"].append(table_info)

        template_info["field_count"] = len(template_info["fields"])
        return template_info

    def extract_template_fields_from_docx(
        self,
        file_path: str
    ) -> List[Dict[str, Any]]:
        """
        从 Word 文档中提取模板字段定义

        适用于比赛评分表格：表格第一列是字段名，第二列是提示词/填写示例

        Args:
            file_path: Word 文件路径

        Returns:
            字段定义列表
        """
        template_info = self.parse_tables_for_template(file_path)

        fields = []
        for field in template_info["fields"]:
            fields.append({
                "cell": f"T{field['table_index']}R{field['row_index']}",  # TableXRowY 格式
                "name": field["field_name"],
                "hint": field["hint"],
                "table_index": field["table_index"],
                "row_index": field["row_index"],
                "field_type": self._infer_field_type_from_hint(field["hint"]),
                "required": True
            })

        return fields

    def _extract_images_info(self, doc: Document, path: Path) -> Dict[str, Any]:
        """
        提取 Word 文档中的图片/嵌入式对象信息

        Args:
            doc: Document 对象
            path: 文件路径

        Returns:
            图片信息字典
        """
        import zipfile
        from io import BytesIO

        image_count = 0
        image_descriptions = []
        inline_shapes_count = 0

        try:
            # 方法1: 通过 inline shapes 统计图片
            try:
                inline_shapes_count = len(doc.inline_shapes)
                if inline_shapes_count > 0:
                    image_count = inline_shapes_count
                    image_descriptions.append(f"文档包含 {inline_shapes_count} 个嵌入式图形/图片")
            except Exception:
                pass

            # 方法2: 通过 ZIP 分析 document.xml 获取图片引用
            try:
                with zipfile.ZipFile(path, 'r') as zf:
                    # 查找 word/media 目录下的图片文件
                    media_files = [f for f in zf.namelist() if f.startswith('word/media/')]
                    if media_files and not inline_shapes_count:
                        image_count = len(media_files)
                        image_descriptions.append(f"文档包含 {image_count} 个嵌入图片")

                    # 检查是否有页眉页脚中的图片
                    header_images = [f for f in zf.namelist() if 'header' in f.lower() and f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
                    if header_images:
                        image_descriptions.append(f"页眉/页脚包含 {len(header_images)} 个图片")
            except Exception:
                pass

        except Exception as e:
            logger.warning(f"提取图片信息失败: {str(e)}")

        return {
            "image_count": image_count,
            "inline_shapes_count": inline_shapes_count,
            "descriptions": image_descriptions,
            "has_images": image_count > 0
        }

    def _infer_field_type_from_hint(self, hint: str) -> str:
        """
        从提示词推断字段类型

        Args:
            hint: 字段提示词

        Returns:
            字段类型 (text/number/date)
        """
        hint_lower = hint.lower()

        # 日期关键词
        date_keywords = ["年", "月", "日", "日期", "时间", "出生"]
        if any(kw in hint for kw in date_keywords):
            return "date"

        # 数字关键词
        number_keywords = ["数量", "金额", "人数", "面积", "增长", "比率", "%", "率"]
        if any(kw in hint_lower for kw in number_keywords):
            return "number"

        return "text"
