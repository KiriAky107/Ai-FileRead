"""
表格模板填写服务

从非结构化文档中检索信息并填写到表格模板
"""
import asyncio
import logging
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

from app.core.database import mongodb
from app.services.llm_service import llm_service
from app.core.document_parser import ParserFactory
from app.services.markdown_ai_service import markdown_ai_service
from app.services.rag_service import rag_service
from app.services.word_ai_service import word_ai_service

logger = logging.getLogger(__name__)


@dataclass
class TemplateField:
    """模板字段"""
    cell: str  # 单元格位置，如 "A1"
    name: str  # 字段名称
    field_type: str = "text"  # 字段类型: text/number/date
    required: bool = True
    hint: str = ""  # 字段提示词


@dataclass
class SourceDocument:
    """源文档"""
    doc_id: str
    filename: str
    doc_type: str
    content: str = ""
    structured_data: Dict[str, Any] = field(default_factory=dict)


@dataclass
class FillResult:
    """填写结果"""
    field: str
    values: List[Any] = None  # 支持多个值
    value: Any = ""  # 保留兼容
    source: str = ""  # 来源文档
    confidence: float = 1.0  # 置信度
    warning: str = None  # 多值提示

    def __post_init__(self):
        if self.values is None:
            self.values = []


class TemplateFillService:
    """表格填写服务"""

    # 通用表头语义扩展字典
    GENERIC_HEADER_EXPANSION = {
        "机构": ["医院", "学校", "企业", "机关", "团体", "协会", "基金会", "研究所", "医院数量", "学校数量", "企业数量"],
        "名称": ["医院名称", "学校名称", "企业名称", "机构名称", "单位名称", "名称"],
        "类型": ["医院类型", "学校类型", "企业类型", "机构类型", "类型分类"],
        "数量": ["医院数量", "学校数量", "企业数量", "机构数量", "个数", "总数", "人员数量"],
        "金额": ["金额", "收入", "支出", "产值", "销售额", "利润", "税收"],
        "比率": ["增长率", "占比", "比重", "比率", "百分比", "使用率", "就业率"],
        "面积": ["占地面积", "建筑面积", "用地面积", "耕地面积", "绿化面积"],
        "人口": ["常住人口", "户籍人口", "流动人口", "城镇人口", "农村人口"],
        "价格": ["价格", "物价", "CPI", "涨幅", "指数"],
        "增长": ["增速", "增长率", "增幅", "增长", "上涨", "下降"],
    }

    # 模板表头到源文档表头的映射缓存
    _header_mapping_cache: Dict[str, Dict[str, str]] = {}

    def _analyze_source_table_structure(self, source_docs: List["SourceDocument"]) -> Dict[str, Any]:
        """
        分析源文档的表格结构

        Args:
            source_docs: 源文档列表

        Returns:
            表格结构分析结果，包含所有表头和样本数据
        """
        table_structures = {}

        for doc_idx, doc in enumerate(source_docs):
            structured = doc.structured_data if doc.structured_data else {}

            # 处理多 sheet 格式
            if structured.get("sheets"):
                for sheet_name, sheet_data in structured.get("sheets", {}).items():
                    if isinstance(sheet_data, dict):
                        columns = sheet_data.get("columns", [])
                        rows = sheet_data.get("rows", [])[:10]  # 只取前10行作为样本
                        key = f"doc{doc_idx}_{sheet_name}"
                        table_structures[key] = {
                            "doc_idx": doc_idx,
                            "sheet_name": sheet_name,
                            "columns": columns,
                            "sample_rows": rows,
                            "column_count": len(columns),
                            "row_count": len(sheet_data.get("rows", []))
                        }

            # 处理 tables 格式
            elif structured.get("tables"):
                for table_idx, table in enumerate(structured.get("tables", [])[:5]):
                    if isinstance(table, dict):
                        headers = table.get("headers", [])
                        rows = table.get("rows", [])[:10]
                        key = f"doc{doc_idx}_table{table_idx}"
                        table_structures[key] = {
                            "doc_idx": doc_idx,
                            "table_idx": table_idx,
                            "columns": headers,
                            "sample_rows": rows,
                            "column_count": len(headers),
                            "row_count": len(table.get("rows", []))
                        }

            # 处理单 sheet 格式
            elif structured.get("columns") and structured.get("rows"):
                columns = structured.get("columns", [])
                rows = structured.get("rows", [])[:10]
                key = f"doc{doc_idx}_default"
                table_structures[key] = {
                    "doc_idx": doc_idx,
                    "columns": columns,
                    "sample_rows": rows,
                    "column_count": len(columns),
                    "row_count": len(structured.get("rows", []))
                }

        logger.info(f"分析源文档表格结构: {len(table_structures)} 个表格")
        return table_structures

    def _build_adaptive_header_mapping(
        self,
        template_fields: List["TemplateField"],
        source_table_structures: Dict[str, Any]
    ) -> Dict[str, Dict[str, Any]]:
        """
        自适应构建模板表头到源文档表头的映射

        Args:
            template_fields: 模板字段列表
            source_table_structures: 源文档表格结构

        Returns:
            映射字典: {field_name: {source_table_key: {column: idx, match_score: score}}}
        """
        mappings = {}

        for field in template_fields:
            field_name = field.name
            field_lower = field_name.lower()
            field_keywords = set(field_lower.replace(" ", "").split())

            best_matches = {}

            for table_key, table_info in source_table_structures.items():
                columns = table_info.get("columns", [])
                if not columns:
                    continue

                best_col_idx = None
                best_col_name = None
                best_score = 0

                for col_idx, col in enumerate(columns):
                    col_str = str(col).strip()
                    col_lower = col_str.lower()
                    col_keywords = set(col_lower.replace(" ", "").split())

                    score = 0

                    # 1. 精确匹配
                    if col_lower == field_lower:
                        score = 1.0
                    # 2. 子字符串匹配
                    elif field_lower in col_lower or col_lower in field_lower:
                        score = 0.8 * max(len(field_lower), len(col_lower)) / min(len(field_lower) + 1, len(col_lower) + 1)
                    # 3. 关键词重叠
                    else:
                        overlap = field_keywords & col_keywords
                        if overlap:
                            score = 0.6 * len(overlap) / max(len(field_keywords), len(col_keywords), 1)

                    # 4. 检查通用表头扩展
                    if score < 0.5:
                        for generic, specifics in self.GENERIC_HEADER_EXPANSION.items():
                            if generic in field_lower:
                                for specific in specifics:
                                    if specific in col_lower or col_lower in specific:
                                        score = 0.7
                                        break
                            if score >= 0.5:
                                break

                    if score > best_score:
                        best_score = score
                        best_col_idx = col_idx
                        best_col_name = col_str

                if best_score >= 0.3 and best_col_idx is not None:
                    best_matches[table_key] = {
                        "column_index": best_col_idx,
                        "column_name": best_col_name,
                        "match_score": best_score,
                        "table_info": table_info
                    }

            if best_matches:
                mappings[field_name] = best_matches
                logger.info(f"字段 '{field_name}' 匹配到 {len(best_matches)} 个源表头，最佳匹配: {list(best_matches.values())[0].get('column_name')}")

        return mappings

    def _extract_with_adaptive_mapping(
        self,
        source_docs: List["SourceDocument"],
        field_name: str,
        mapping: Dict[str, Dict[str, Any]]
    ) -> List[str]:
        """
        使用自适应映射提取字段值

        Args:
            source_docs: 源文档列表
            field_name: 字段名
            mapping: 字段到源表头的映射

        Returns:
            提取的值列表
        """
        values = []

        if field_name not in mapping:
            return values

        best_matches = mapping[field_name]

        for table_key, match_info in best_matches.items():
            table_info = match_info.get("table_info", {})
            col_idx = match_info.get("column_index", 0)
            doc_idx = table_info.get("doc_idx", 0)

            if doc_idx >= len(source_docs):
                continue

            doc = source_docs[doc_idx]
            structured = doc.structured_data if doc.structured_data else {}

            # 根据表格类型提取值
            rows = []

            # 多 sheet 格式
            if structured.get("sheets"):
                sheet_name = table_info.get("sheet_name")
                if sheet_name:
                    sheet_data = structured.get("sheets", {}).get(sheet_name, {})
                    rows = sheet_data.get("rows", [])

            # tables 格式
            elif structured.get("tables"):
                table_idx = table_info.get("table_idx", 0)
                tables = structured.get("tables", [])
                if table_idx < len(tables):
                    rows = tables[table_idx].get("rows", [])

            # 单 sheet 格式
            elif structured.get("rows"):
                rows = structured.get("rows", [])

            # 提取指定列的值
            for row in rows:
                if isinstance(row, list) and col_idx < len(row):
                    val = self._format_value(row[col_idx])
                    if val and self._is_valid_data_value(val):
                        values.append(val)
                elif isinstance(row, dict):
                    # 对于 dict 格式的行
                    columns = table_info.get("columns", [])
                    if col_idx < len(columns):
                        col_name = columns[col_idx]
                        val = self._format_value(row.get(col_name, ""))
                        if val and self._is_valid_data_value(val):
                            values.append(val)

        # 过滤和去重
        seen = set()
        unique_values = []
        for v in values:
            if v not in seen:
                seen.add(v)
                unique_values.append(v)

        return unique_values

    def __init__(self):
        self.llm = llm_service

    async def fill_template(
        self,
        template_fields: List[TemplateField],
        source_doc_ids: Optional[List[str]] = None,
        source_file_paths: Optional[List[str]] = None,
        user_hint: Optional[str] = None,
        template_id: Optional[str] = None,
        template_file_type: Optional[str] = "xlsx",
        task_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        填写表格模板

        Args:
            template_fields: 模板字段列表
            source_doc_ids: 源文档 MongoDB ID 列表
            source_file_paths: 源文档文件路径列表
            user_hint: 用户提示（如"请从合同文档中提取"）
            template_id: 模板文件路径（用于重新生成表头）
            template_file_type: 模板文件类型
            task_id: 可选的任务ID，用于任务进度跟踪

        Returns:
            填写结果
        """
        filled_data = {}
        fill_details = []

        logger.info(f"开始填表: {len(template_fields)} 个字段, {len(source_doc_ids or [])} 个源文档")
        logger.info(f"source_doc_ids: {source_doc_ids}")
        logger.info(f"source_file_paths: {source_file_paths}")

        # 1. 加载源文档内容
        source_docs = await self._load_source_documents(source_doc_ids, source_file_paths)

        logger.info(f"加载了 {len(source_docs)} 个源文档")

        # 打印每个加载的文档的详细信息
        for i, doc in enumerate(source_docs):
            logger.info(f"  文档[{i}]: id={doc.doc_id}, filename={doc.filename}, doc_type={doc.doc_type}")
            logger.info(f"    content长度: {len(doc.content)}, structured_data keys: {list(doc.structured_data.keys()) if doc.structured_data else 'None'}")

        if not source_docs:
            logger.warning("没有找到源文档，填表结果将全部为空")

        # 3. 检查是否需要使用源文档重新生成表头
        # 条件：源文档已加载 AND 现有字段看起来是自动生成的（如"字段1"、"字段2"）
        needs_regenerate_headers = (
            len(source_docs) > 0 and
            len(template_fields) > 0 and
            all(self._is_auto_generated_field(f.name) for f in template_fields)
        )

        if needs_regenerate_headers:
            logger.info(f"检测到自动生成表头，尝试使用源文档重新生成... (当前字段: {[f.name for f in template_fields]})")

            # 将 SourceDocument 转换为 source_contents 格式
            source_contents = []
            for doc in source_docs:
                structured = doc.structured_data if doc.structured_data else {}

                # 获取标题
                titles = structured.get("titles", [])
                if not titles:
                    titles = []

                # 获取表格
                tables = structured.get("tables", [])
                tables_count = len(tables) if tables else 0

                # 生成表格摘要
                tables_summary = ""
                if tables:
                    tables_summary = "\n【文档中的表格】:\n"
                    for idx, table in enumerate(tables[:5]):
                        if isinstance(table, dict):
                            headers = table.get("headers", [])
                            rows = table.get("rows", [])
                            if headers:
                                tables_summary += f"表格{idx+1}表头: {', '.join(str(h) for h in headers)}\n"
                            if rows:
                                tables_summary += f"表格{idx+1}前3行: "
                                for row_idx, row in enumerate(rows[:3]):
                                    if isinstance(row, list):
                                        tables_summary += " | ".join(str(c) for c in row) + "; "
                                    elif isinstance(row, dict):
                                        tables_summary += " | ".join(str(row.get(h, "")) for h in headers if headers) + "; "
                                tables_summary += "\n"

                source_contents.append({
                    "filename": doc.filename,
                    "doc_type": doc.doc_type,
                    "content": doc.content[:5000] if doc.content else "",
                    "titles": titles[:10] if titles else [],
                    "tables_count": tables_count,
                    "tables_summary": tables_summary
                })

            # 使用源文档内容重新生成表头
            if template_id and template_file_type:
                logger.info(f"使用源文档重新生成表头: template_id={template_id}, template_file_type={template_file_type}")
                new_fields = await self.get_template_fields_from_file(
                    template_id,
                    template_file_type,
                    source_contents=source_contents
                )
                if new_fields and len(new_fields) > 0:
                    logger.info(f"成功重新生成表头: {[f.name for f in new_fields]}")
                    template_fields = new_fields
                else:
                    logger.warning("重新生成表头返回空结果，使用原始字段")
            else:
                logger.warning("无法重新生成表头：缺少 template_id 或 template_file_type")
        else:
            if source_docs and template_fields:
                logger.info(f"表头看起来正常（非自动生成），无需重新生成: {[f.name for f in template_fields[:5]]}")

        # 2. 并行提取所有字段（跳过AI审核以提升速度）
        logger.info(f"开始并行提取 {len(template_fields)} 个字段...")

        # 并行处理所有字段
        tasks = []
        for idx, field in enumerate(template_fields):
            task = self._extract_single_field_fast(
                field=field,
                source_docs=source_docs,
                user_hint=user_hint,
                field_idx=idx,
                total_fields=len(template_fields)
            )
            tasks.append(task)

        # 等待所有任务完成
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 处理结果
        for idx, result in enumerate(results):
            field = template_fields[idx]
            if isinstance(result, Exception):
                logger.error(f"提取字段 {field.name} 失败: {str(result)}")
                filled_data[field.name] = [f"[提取失败: {str(result)}]"]
                fill_details.append({
                    "field": field.name,
                    "cell": field.cell,
                    "values": [f"[提取失败]"],
                    "value": f"[提取失败]",
                    "source": "error",
                    "confidence": 0.0
                })
            else:
                filled_data[field.name] = result.values if result.values else [""]
                fill_details.append({
                    "field": field.name,
                    "cell": field.cell,
                    "values": result.values,
                    "value": result.value,
                    "source": result.source,
                    "confidence": result.confidence,
                    "warning": result.warning
                })
                logger.info(f"字段 {field.name} 填写完成: {len(result.values) if result.values else 0} 个值")

        # 计算最大行数
        max_rows = max(len(v) for v in filled_data.values()) if filled_data else 1
        logger.info(f"填表完成: {len(filled_data)} 个字段, 最大行数: {max_rows}")

        return {
            "success": True,
            "filled_data": filled_data,
            "fill_details": fill_details,
            "source_doc_count": len(source_docs),
            "max_rows": max_rows
        }

    async def _load_source_documents(
        self,
        source_doc_ids: Optional[List[str]] = None,
        source_file_paths: Optional[List[str]] = None
    ) -> List[SourceDocument]:
        """
        加载源文档内容

        Args:
            source_doc_ids: MongoDB 文档 ID 列表
            source_file_paths: 源文档文件路径列表

        Returns:
            源文档列表
        """
        source_docs = []

        # 1. 从 MongoDB 加载文档
        if source_doc_ids:
            for doc_id in source_doc_ids:
                try:
                    doc = await mongodb.get_document(doc_id)
                    if doc:
                        sd = doc.get("structured_data", {})
                        sd_keys = list(sd.keys()) if sd else []
                        logger.info(f"从MongoDB加载文档: {doc_id}, doc_type={doc.get('doc_type')}, structured_data keys={sd_keys}")

                        # 如果 structured_data 为空，但有 file_path，尝试重新解析文件
                        doc_content = doc.get("content", "")
                        if not sd or (not sd.get("tables") and not sd.get("headers") and not sd.get("rows")):
                            file_path = doc.get("metadata", {}).get("file_path")
                            if file_path:
                                logger.info(f"  structured_data 为空，尝试重新解析文件: {file_path}")
                                try:
                                    parser = ParserFactory.get_parser(file_path)
                                    result = parser.parse(file_path)
                                    if result.success and result.data:
                                        if result.data.get("structured_data"):
                                            sd = result.data.get("structured_data")
                                            logger.info(f"  重新解析成功，structured_data keys: {list(sd.keys())}")
                                        elif result.data.get("tables"):
                                            sd = {"tables": result.data.get("tables", [])}
                                            logger.info(f"  使用 data.tables，tables数量: {len(sd.get('tables', []))}")
                                        elif result.data.get("rows"):
                                            sd = result.data
                                            logger.info(f"  使用 data.rows 格式")
                                        if result.data.get("content"):
                                            doc_content = result.data.get("content", "")
                                    else:
                                        logger.warning(f"  重新解析失败: {result.error if result else 'unknown'}")
                                except Exception as parse_err:
                                    logger.error(f"  重新解析文件异常: {str(parse_err)}")

                        if sd.get("tables"):
                            logger.info(f"  tables数量: {len(sd.get('tables', []))}")
                            if sd["tables"]:
                                first_table = sd["tables"][0]
                                logger.info(f"  第一表格: headers={first_table.get('headers', [])[:3]}..., rows数量={len(first_table.get('rows', []))}")

                        source_docs.append(SourceDocument(
                            doc_id=doc_id,
                            filename=doc.get("metadata", {}).get("original_filename", "unknown"),
                            doc_type=doc.get("doc_type", "unknown"),
                            content=doc_content,
                            structured_data=sd
                        ))
                except Exception as e:
                    logger.error(f"从MongoDB加载文档失败 {doc_id}: {str(e)}")

        # 2. 从文件路径加载文档
        if source_file_paths:
            for file_path in source_file_paths:
                try:
                    file_ext = file_path.lower().split('.')[-1]

                    # 对于 Word 文档，优先使用 AI 解析
                    if file_ext == 'docx':
                        # 使用 AI 深度解析 Word 文档
                        ai_result = await word_ai_service.parse_word_with_ai(
                            file_path=file_path,
                            user_hint="请提取文档中的所有结构化数据，包括表格、键值对等"
                        )

                        if ai_result.get("success"):
                            # AI 解析成功，转换为 SourceDocument 格式
                            parse_type = ai_result.get("type", "unknown")

                            # 构建 structured_data
                            doc_structured = {
                                "ai_parsed": True,
                                "parse_type": parse_type,
                                "tables": [],
                                "key_values": ai_result.get("key_values", {}) if "key_values" in ai_result else {},
                                "list_items": ai_result.get("list_items", []) if "list_items" in ai_result else [],
                                "summary": ai_result.get("summary", "") if "summary" in ai_result else ""
                            }

                            # 如果 AI 返回了表格数据
                            if parse_type == "table_data":
                                headers = ai_result.get("headers", [])
                                rows = ai_result.get("rows", [])
                                if headers and rows:
                                    doc_structured["tables"] = [{
                                        "headers": headers,
                                        "rows": rows
                                    }]
                                    doc_structured["columns"] = headers
                                    doc_structured["rows"] = rows
                                    logger.info(f"AI 表格数据: {len(headers)} 列, {len(rows)} 行")
                            elif parse_type == "structured_text":
                                tables = ai_result.get("tables", [])
                                if tables:
                                    doc_structured["tables"] = tables
                                    logger.info(f"AI 结构化文本提取到 {len(tables)} 个表格")

                            # 获取摘要内容
                            content_text = doc_structured.get("summary", "") or ai_result.get("description", "")

                            source_docs.append(SourceDocument(
                                doc_id=file_path,
                                filename=file_path.split("/")[-1] if "/" in file_path else file_path.split("\\")[-1],
                                doc_type="docx",
                                content=content_text,
                                structured_data=doc_structured
                            ))
                            logger.info(f"AI 解析 Word 文档: {file_path}, type={parse_type}, tables={len(doc_structured.get('tables', []))}")
                            continue  # 跳后续的基础解析

                    # 基础解析（Excel 或非 AI 解析的 Word）
                    parser = ParserFactory.get_parser(file_path)
                    result = parser.parse(file_path)
                    if result.success:
                        # result.data 的结构取决于解析器类型:
                        # - Excel 单 sheet: {columns: [...], rows: [...], row_count, column_count}
                        # - Excel 多 sheet: {sheets: {sheet_name: {columns, rows, ...}}}
                        # - Markdown: {content: "...", tables: [...], structured_data: {tables: [...]}}
                        # - Word/TXT: {content: "...", structured_data: {...}}
                        doc_data = result.data if result.data else {}
                        doc_content = doc_data.get("content", "") if isinstance(doc_data, dict) else ""

                        # 检查并提取 structured_data
                        doc_structured = {}
                        if isinstance(doc_data, dict):
                            logger.info(f"文档 {file_path} doc_data keys: {list(doc_data.keys())}")

                            # Excel 多 sheet
                            if "sheets" in doc_data:
                                doc_structured = doc_data
                                logger.info(f"  -> 使用 Excel 多 sheet 格式")
                            # Excel 单 sheet 或有 rows 的格式
                            elif "rows" in doc_data:
                                doc_structured = doc_data
                                logger.info(f"  -> 使用 rows 格式，列数: {len(doc_data.get('columns', []))}")
                            # Markdown 格式：tables 可能直接在 doc_data.tables 或在 structured_data.tables 中
                            elif "tables" in doc_data and doc_data["tables"]:
                                # Markdown: tables 直接在 doc_data 中
                                tables = doc_data["tables"]
                                first_table = tables[0]
                                doc_structured = {
                                    "headers": first_table.get("headers", []),
                                    "rows": first_table.get("rows", [])
                                }
                                logger.info(f"  -> 使用 doc_data.tables 格式，表头: {doc_structured.get('headers', [])[:5]}")
                            elif "structured_data" in doc_data and isinstance(doc_data["structured_data"], dict):
                                # Markdown: tables 在 structured_data 中
                                tables = doc_data["structured_data"].get("tables", [])
                                if tables:
                                    first_table = tables[0]
                                    doc_structured = {
                                        "headers": first_table.get("headers", []),
                                        "rows": first_table.get("rows", [])
                                    }
                                    logger.info(f"  -> 使用 structured_data.tables 格式，表头: {doc_structured.get('headers', [])[:5]}")
                                else:
                                    logger.warning(f"  -> structured_data.tables 为空")
                            else:
                                logger.warning(f"  -> 未识别的文档格式，无 structured_data")

                        source_docs.append(SourceDocument(
                            doc_id=file_path,
                            filename=result.metadata.get("filename", file_path.split("/")[-1]),
                            doc_type=result.metadata.get("extension", "unknown").replace(".", ""),
                            content=doc_content,
                            structured_data=doc_structured
                        ))
                        logger.info(f"从文件加载文档: {file_path}, content长度: {len(doc_content)}, structured数据: {bool(doc_structured)}")
                except Exception as e:
                    logger.error(f"从文件加载文档失败 {file_path}: {str(e)}")

        return source_docs

    async def _extract_field_value(
        self,
        field: TemplateField,
        source_docs: List[SourceDocument],
        user_hint: Optional[str] = None
    ) -> FillResult:
        """
        使用 LLM 从源文档中提取字段值

        Args:
            field: 字段定义
            source_docs: 源文档列表
            user_hint: 用户提示

        Returns:
            提取结果
        """
        if not source_docs:
            return FillResult(
                field=field.name,
                value="",
                source="无源文档",
                confidence=0.0
            )

        # 优先尝试直接从结构化数据中提取列值（适用于 Excel 等有 rows 的数据）
        direct_values = self._extract_values_from_structured_data(source_docs, field.name)
        if direct_values:
            logger.info(f"✅ 字段 {field.name} 直接从结构化数据提取到 {len(direct_values)} 个值")
            return FillResult(
                field=field.name,
                values=direct_values,
                value=direct_values[0] if direct_values else "",
                source="结构化数据直接提取",
                confidence=1.0
            )

        # 无法直接从结构化数据提取，尝试 AI 分析非结构化文档
        ai_structured = await self._analyze_unstructured_docs_for_fields(source_docs, field, user_hint)
        if ai_structured:
            logger.info(f"✅ 字段 {field.name} 通过 AI 分析结构化提取到数据")
            return ai_structured

        # 无法从结构化数据提取，使用 LLM
        logger.info(f"字段 {field.name} 无法直接从结构化数据提取，使用 LLM...")

        # 构建上下文文本 - 传入字段名，只提取该列数据
        context_text = self._build_context_text(source_docs, field_name=field.name, max_length=200000)

        # 构建提示词
        hint_text = field.hint if field.hint else f"请提取{field.name}的信息"
        if user_hint:
            hint_text = f"{user_hint}。{hint_text}"

        prompt = f"""你是一个专业的数据提取专家。请从以下文档内容中提取与"{field.name}"相关的所有信息。

提示词: {hint_text}

文档内容：
{context_text}

请分析文档结构（可能包含表格、标题段落等），找出所有与"{field.name}"相关的数据。
如果找到表格数据，返回多行值；如果是非表格段落，提取关键信息。

请严格按照以下 JSON 格式输出：
{{
    "values": ["第1行的值", "第2行的值", ...],
    "source": "数据来源描述",
    "confidence": 0.0到1.0之间的置信度
}}
"""

        # 调用 LLM
        messages = [
            {"role": "system", "content": "你是一个专业的数据提取助手。请严格按JSON格式输出。"},
            {"role": "user", "content": prompt}
        ]

        try:
            response = await self.llm.chat(
                messages=messages,
                temperature=0.1,
                max_tokens=4000
            )

            content = self.llm.extract_message_content(response)

            # 解析 JSON 响应
            import json
            import re

            extracted_values = []
            extracted_value = ""
            extracted_source = "LLM生成"
            confidence = 0.5

            logger.info(f"原始 LLM 返回: {content[:500]}")

            # ========== 步骤1: 彻底清理 markdown 和各种格式问题 ==========
            # 移除 ```json 和 ``` 标记
            cleaned = content.strip()
            cleaned = re.sub(r'^```json\s*', '', cleaned, flags=re.MULTILINE)
            cleaned = re.sub(r'^```\s*', '', cleaned, flags=re.MULTILINE)
            cleaned = cleaned.strip()

            logger.info(f"清理后: {cleaned[:500]}")

            # ========== 步骤2: 定位 JSON 开始位置 ==========
            json_start = -1
            # 找到第一个 { 或 [
            for i, c in enumerate(cleaned):
                if c == '{' or c == '[':
                    json_start = i
                    break

            if json_start == -1:
                logger.warning(f"无法找到 JSON 开始位置")
                extracted_values = self._extract_values_from_text(cleaned, field.name)
            else:
                json_text = cleaned[json_start:]
                logger.info(f"JSON 开始位置: {json_start}, 内容: {json_text[:200]}")

                # ========== 步骤3: 尝试解析 JSON ==========
                # 3a. 尝试直接解析整个字符串
                try:
                    result = json.loads(json_text)
                    extracted_values = self._extract_values_from_json(result)
                    if extracted_values:
                        logger.info(f"✅ 直接解析成功，得到 {len(extracted_values)} 个值")
                    else:
                        logger.warning(f"直接解析成功但未提取到值")
                except json.JSONDecodeError as e:
                    logger.warning(f"直接解析失败: {e}, 尝试修复...")

                    # 3b. 尝试修复常见的 JSON 问题
                    # 尝试1: 找到配对的闭合括号
                    fixed_json = self._fix_json(json_text)
                    if fixed_json:
                        try:
                            result = json.loads(fixed_json)
                            extracted_values = self._extract_values_from_json(result)
                            if extracted_values:
                                logger.info(f"✅ 修复后解析成功，得到 {len(extracted_values)} 个值")
                        except json.JSONDecodeError as e2:
                            logger.warning(f"修复后仍然失败: {e2}")

                    # 3c. 如果以上都失败，使用正则直接从文本提取 values 数组
                    if not extracted_values:
                        extracted_values = self._extract_values_by_regex(cleaned)
                        if extracted_values:
                            logger.info(f"✅ 正则提取成功，得到 {len(extracted_values)} 个值")
                        else:
                            # 最后的备选：使用旧的文本提取
                            extracted_values = self._extract_values_from_text(cleaned, field.name)

            # 如果仍然没有提取到值
            if not extracted_values:
                extracted_values = [""]
                logger.warning(f"❌ 字段 {field.name} 没有提取到值")

            logger.info(f"✅✅ 字段 {field.name} 最终返回: {len(extracted_values)} 个值, 示例: {extracted_values[:3]}")

            return FillResult(
                field=field.name,
                values=extracted_values,
                value=extracted_values[0] if extracted_values else "",
                source=extracted_source,
                confidence=confidence
            )

        except Exception as e:
            logger.error(f"LLM 提取失败: {str(e)}")
            return FillResult(
                field=field.name,
                values=[""],
                value="",
                source=f"提取失败: {str(e)}",
                confidence=0.0
            )

    async def _extract_single_field_fast(
        self,
        field: TemplateField,
        source_docs: List[SourceDocument],
        user_hint: Optional[str] = None,
        field_idx: int = 0,
        total_fields: int = 1
    ) -> FillResult:
        """
        快速提取单个字段（跳过AI审核，减少LLM调用）

        Args:
            field: 字段定义
            source_docs: 源文档列表
            user_hint: 用户提示
            field_idx: 当前字段索引（用于日志）
            total_fields: 总字段数（用于日志）

        Returns:
            提取结果
        """
        try:
            if not source_docs:
                return FillResult(
                    field=field.name,
                    value="",
                    values=[""],
                    source="无源文档",
                    confidence=0.0
                )

            # 1. 优先尝试直接从结构化数据中提取（最快路径）
            direct_values = self._extract_values_from_structured_data(source_docs, field.name)
            if direct_values:
                logger.info(f"✅ [{field_idx+1}/{total_fields}] 字段 {field.name} 直接从结构化数据提取到 {len(direct_values)} 个值")
                return FillResult(
                    field=field.name,
                    values=direct_values,
                    value=direct_values[0] if direct_values else "",
                    source="结构化数据直接提取",
                    confidence=1.0
                )

            # 2. 无法直接从结构化数据提取，使用简化版AI提取
            logger.info(f"🔍 [{field_idx+1}/{total_fields}] 字段 {field.name} 尝试AI提取...")

            # 构建提示词 - 简化版
            hint_text = field.hint if field.hint else f"请提取{field.name}的信息"
            if user_hint:
                hint_text = f"{user_hint}。{hint_text}"

            # 优先使用 RAG 检索内容，否则使用文档开头部分
            context_parts = []
            for doc in source_docs:
                if not doc.content:
                    logger.info(f"  文档 {doc.filename} 无content内容")
                    continue

                logger.info(f"  处理文档: {doc.filename}, doc_id={doc.doc_id}, content长度={len(doc.content)}")

                # 尝试 RAG 检索
                rag_results = rag_service.retrieve(
                    query=f"{field.name} {hint_text}",
                    top_k=3,
                    min_score=0.1
                )

                if rag_results:
                    logger.info(f"  RAG检索到 {len(rag_results)} 条结果")
                    # 使用 RAG 检索到的内容
                    for r in rag_results:
                        rag_doc_id = r.get("doc_id", "")
                        if rag_doc_id.startswith(doc.doc_id):
                            context_parts.append(r["content"])
                            logger.info(f"    匹配成功，使用RAG内容，长度={len(r['content'])}")
                else:
                    # RAG 没结果，使用文档内容开头
                    context_parts.append(doc.content[:2500])
                    logger.info(f"  RAG无结果，使用文档开头 {min(2500, len(doc.content))} 字符")

            context = "\n\n".join(context_parts[:3]) if context_parts else ""
            logger.info(f"  最终context长度: {len(context)}, 内容预览: {context[:200] if context else '无'}...")

            prompt = f"""你是一个专业的数据提取专家。请严格按照表头字段「{field.name}」从文档中提取数据。

提示: {hint_text}

【重要规则 - 必须遵守】
1. **每个值必须有标注**：根据数据来源添加合适的标注前缀！
   - ✅ 正确格式：
     - "2024年：38710个"
     - "北京：1234万人次"
     - "某省：5678万人"
     - "公立医院：11754个"
     - "三级医院：4111个"
     - "图书馆：3246个"
   - ❌ 错误格式："38710个"（缺少标注）

2. **标注类型根据数据决定**：
   - 年份类数据 → "2024年：xxx"、"2023年：xxx"
   - 地区类数据 → "北京：xxx"、"广东：xxx"、"某县：xxx"
   - 机构/分类数据 → "公立医院：xxx"、"三级医院：xxx"、"图书馆：xxx"
   - 其他分类 → 根据实际情况标注

3. **严格按表头提取**：只提取与「{field.name}」直接相关的数据

4. **多值必须全部提取并标注**：如果文档中提到多个相关数据，每个都要有标注

文档内容：
{context if context else "（无文档内容）"}

请严格按格式返回JSON：{{"values": ["标注：数值", "标注：数值", ...]}}
注意：values数组中每个元素都必须包含标注前缀，不能只有数值！
"""

            messages = [
                {"role": "system", "content": "你是一个专业的数据提取助手，擅长从政府统计公报中提取数据。严格按JSON格式输出，只返回values数组。"},
                {"role": "user", "content": prompt}
            ]

            response = await self.llm.chat(
                messages=messages,
                temperature=0.1,
                max_tokens=1000
            )

            content = self.llm.extract_message_content(response)
            logger.info(f"  LLM原始返回: {content[:500]}")

            # 解析JSON
            import json
            import re
            cleaned = content.strip()

            # 查找JSON开始位置
            json_start = -1
            for i, c in enumerate(cleaned):
                if c == '{':
                    json_start = i
                    break

            values = []
            source = "AI提取"
            if json_start >= 0:
                try:
                    json_text = cleaned[json_start:]
                    result = json.loads(json_text)
                    values = result.get("values", [])
                    logger.info(f"  JSON解析成功，values: {values}")
                except json.JSONDecodeError as e:
                    logger.warning(f"  JSON解析失败: {e}，尝试修复...")
                    # 尝试修复常见JSON问题
                    try:
                        # 尝试找到values数组
                        values_match = re.search(r'"values"\s*:\s*\[(.*?)\]', cleaned, re.DOTALL)
                        if values_match:
                            values_str = values_match.group(1)
                            # 提取数组中的字符串
                            values = re.findall(r'"([^"]*)"', values_str)
                            logger.info(f"  正则提取values: {values}")
                    except:
                        pass

            # 如果values为空，尝试从文本中用正则提取数字+单位
            if not values or values == [""]:
                logger.info(f"  JSON解析未获取到值，尝试正则提取...")
                # 匹配数字+单位或百分号的模式
                patterns = [
                    r'(\d+\.?\d*[亿万千百十个]?[%‰℃℃万元亿]?)',  # 通用数字+单位
                    r'(\d+\.?\d*%)',  # 百分号
                    r'(\d+\.?\d*[个万人亿元]?)',  # 中文单位
                ]
                for pattern in patterns:
                    matches = re.findall(pattern, context)
                    if matches:
                        # 过滤掉纯数字
                        filtered = [m for m in matches if not m.replace('.', '').isdigit()]
                        if filtered:
                            values = filtered[:10]  # 最多取10个
                            logger.info(f"  正则提取到: {values}")
                            break

            if not values or values == [""]:
                values = self._extract_values_by_regex(cleaned)

            if not values:
                values = [""]

            # 生成多值提示（基于实际检测到的值数量）
            warning = ""
            if len(values) > 1:
                warning = f"⚠️ 检测到 {len(values)} 个值：{values[:5]}{'...' if len(values) > 5 else ''}"

            logger.info(f"✅ [{field_idx+1}/{total_fields}] 字段 {field.name} AI提取完成: {len(values)} 个值")
            if warning:
                logger.info(f"   {warning}")

            return FillResult(
                field=field.name,
                values=values,
                value=values[0] if values else "",
                source=source,
                confidence=0.8,
                warning=warning if warning else None
            )

        except Exception as e:
            logger.error(f"❌ [{field_idx+1}/{total_fields}] 字段 {field.name} 提取失败: {str(e)}")
            return FillResult(
                field=field.name,
                values=[""],
                value="",
                source=f"提取失败: {str(e)}",
                confidence=0.0
            )

    async def _verify_field_value(
        self,
        field: TemplateField,
        extracted_values: List[str],
        source_docs: List[SourceDocument],
        user_hint: Optional[str] = None
    ) -> Optional[FillResult]:
        """
        验证并修正提取的字段值

        Args:
            field: 字段定义
            extracted_values: 已提取的值
            source_docs: 源文档列表
            user_hint: 用户提示

        Returns:
            验证后的结果，如果验证通过返回None（使用原结果）
        """
        if not extracted_values or not extracted_values[0]:
            return None

        if not source_docs:
            return None

        try:
            # 构建验证上下文
            context_text = self._build_context_text(source_docs, field_name=field.name, max_length=15000)

            hint_text = field.hint if field.hint else f"请理解{field.name}字段的含义"
            if user_hint:
                hint_text = f"{user_hint}。{hint_text}"

            prompt = f"""你是一个数据质量审核专家。请审核以下提取的数据是否合理。

【待审核字段】
字段名：{field.name}
字段说明：{hint_text}

【已提取的值】
{extracted_values[:10]}  # 最多审核前10个值

【源文档上下文】
{context_text[:8000]}

【审核要求】
1. 这些值是否符合字段的含义？
2. 值在原文中的原始含义是什么？检查是否有误解或误提取
3. 是否存在明显错误、空值或不合理的数据？
4. 如果表格有多个列，请确认提取的是正确的列

请严格按照以下 JSON 格式输出（只需输出 JSON，不要其他内容）：
{{
    "is_valid": true或false,
    "corrected_values": ["修正后的值列表"] 或 null（如果无需修正）,
    "reason": "审核说明，解释判断理由",
    "original_meaning": "值在原文中的原始含义描述"
}}
"""

            messages = [
                {"role": "system", "content": "你是一个严格的数据质量审核专家。请仔细核对原文和提取的值是否匹配。"},
                {"role": "user", "content": prompt}
            ]

            response = await self.llm.chat(
                messages=messages,
                temperature=0.2,
                max_tokens=3000
            )

            content = self.llm.extract_message_content(response)
            logger.info(f"字段 {field.name} 审核返回: {content[:300]}")

            # 解析 JSON
            import json
            import re

            cleaned = content.strip()
            cleaned = re.sub(r'^```json\s*', '', cleaned, flags=re.MULTILINE)
            cleaned = re.sub(r'^```\s*', '', cleaned, flags=re.MULTILINE)
            cleaned = cleaned.strip()

            json_start = -1
            for i, c in enumerate(cleaned):
                if c == '{':
                    json_start = i
                    break

            if json_start == -1:
                logger.warning(f"字段 {field.name} 审核：无法找到 JSON")
                return None

            json_text = cleaned[json_start:]
            result = json.loads(json_text)

            is_valid = result.get("is_valid", True)
            corrected_values = result.get("corrected_values")
            reason = result.get("reason", "")
            original_meaning = result.get("original_meaning", "")

            logger.info(f"字段 {field.name} 审核结果: is_valid={is_valid}, reason={reason[:100]}")

            if not is_valid and corrected_values:
                # 值有问题且有修正建议，使用修正后的值
                logger.info(f"字段 {field.name} 使用修正后的值: {corrected_values[:5]}")
                return FillResult(
                    field=field.name,
                    values=corrected_values,
                    value=corrected_values[0] if corrected_values else "",
                    source=f"AI审核修正: {reason[:100]}",
                    confidence=0.7
                )
            elif not is_valid and original_meaning:
                # 值有问题但无修正，记录原始含义供用户参考
                logger.info(f"字段 {field.name} 审核发现问题: {original_meaning}")
                return FillResult(
                    field=field.name,
                    values=extracted_values,
                    value=extracted_values[0] if extracted_values else "",
                    source=f"AI审核疑问: {original_meaning[:100]}",
                    confidence=0.5
                )

            # 验证通过，返回 None 表示使用原结果
            return None

        except Exception as e:
            logger.error(f"字段 {field.name} 审核失败: {str(e)}")
            return None

    def _build_context_text(self, source_docs: List[SourceDocument], field_name: str = None, max_length: int = 8000) -> str:
        """
        构建上下文文本

        Args:
            source_docs: 源文档列表
            field_name: 需要提取的字段名（可选，用于只提取特定列）
            max_length: 最大字符数

        Returns:
            上下文文本
        """
        contexts = []
        total_length = 0

        for doc in source_docs:
            # 优先使用结构化数据（表格），其次使用文本内容
            doc_content = ""
            row_count = 0

            if doc.structured_data and doc.structured_data.get("sheets"):
                # parse_all_sheets 格式: {sheets: {sheet_name: {columns, rows}}}
                sheets = doc.structured_data.get("sheets", {})
                for sheet_name, sheet_data in sheets.items():
                    if isinstance(sheet_data, dict):
                        columns = sheet_data.get("columns", [])
                        rows = sheet_data.get("rows", [])
                        if rows and columns:
                            doc_content += f"\n【文档: {doc.filename} - {sheet_name}，共 {len(rows)} 行】\n"
                            # 如果指定了字段名，只提取该列数据
                            if field_name:
                                # 查找匹配的列（模糊匹配）
                                target_col = None
                                for col in columns:
                                    if field_name.lower() in str(col).lower() or str(col).lower() in field_name.lower():
                                        target_col = col
                                        break
                                if target_col:
                                    doc_content += f"列名: {target_col}\n"
                                    for row_idx, row in enumerate(rows):
                                        if isinstance(row, dict):
                                            val = row.get(target_col, "")
                                        elif isinstance(row, list) and target_col in columns:
                                            val = row[columns.index(target_col)]
                                        else:
                                            val = ""
                                        doc_content += f"行{row_idx+1}: {val}\n"
                                        row_count += 1
                                else:
                                    # 列名不匹配，输出所有列（但只输出关键列）
                                    doc_content += " | ".join(str(col) for col in columns) + "\n"
                                    for row in rows:
                                        if isinstance(row, dict):
                                            doc_content += " | ".join(str(row.get(col, "")) for col in columns) + "\n"
                                        elif isinstance(row, list):
                                            doc_content += " | ".join(str(cell) for cell in row) + "\n"
                                        row_count += 1
                            else:
                                # 输出所有列和行
                                doc_content += " | ".join(str(col) for col in columns) + "\n"
                                for row in rows:
                                    if isinstance(row, dict):
                                        doc_content += " | ".join(str(row.get(col, "")) for col in columns) + "\n"
                                    elif isinstance(row, list):
                                        doc_content += " | ".join(str(cell) for cell in row) + "\n"
                                    row_count += 1
            elif doc.structured_data and doc.structured_data.get("rows"):
                # Excel 单 sheet 格式: {columns: [...], rows: [...], ...}
                columns = doc.structured_data.get("columns", [])
                rows = doc.structured_data.get("rows", [])
                if rows and columns:
                    doc_content += f"\n【文档: {doc.filename}，共 {len(rows)} 行】\n"
                    if field_name:
                        target_col = None
                        for col in columns:
                            if field_name.lower() in str(col).lower() or str(col).lower() in field_name.lower():
                                target_col = col
                                break
                        if target_col:
                            doc_content += f"列名: {target_col}\n"
                            for row_idx, row in enumerate(rows):
                                if isinstance(row, dict):
                                    val = row.get(target_col, "")
                                elif isinstance(row, list) and target_col in columns:
                                    val = row[columns.index(target_col)]
                                else:
                                    val = ""
                                doc_content += f"行{row_idx+1}: {val}\n"
                                row_count += 1
                        else:
                            doc_content += " | ".join(str(col) for col in columns) + "\n"
                            for row in rows:
                                if isinstance(row, dict):
                                    doc_content += " | ".join(str(row.get(col, "")) for col in columns) + "\n"
                                elif isinstance(row, list):
                                    doc_content += " | ".join(str(cell) for cell in row) + "\n"
                                row_count += 1
                    else:
                        doc_content += " | ".join(str(col) for col in columns) + "\n"
                        for row in rows:
                            if isinstance(row, dict):
                                doc_content += " | ".join(str(row.get(col, "")) for col in columns) + "\n"
                            elif isinstance(row, list):
                                doc_content += " | ".join(str(cell) for cell in row) + "\n"
                            row_count += 1
            elif doc.structured_data and doc.structured_data.get("tables"):
                # Markdown 表格格式: {tables: [{headers: [...], rows: [...]}]}
                tables = doc.structured_data.get("tables", [])
                for table in tables:
                    if isinstance(table, dict):
                        headers = table.get("headers", [])
                        rows = table.get("rows", [])
                        if rows and headers:
                            doc_content += f"\n【文档: {doc.filename} - 表格】\n"
                            doc_content += " | ".join(str(h) for h in headers) + "\n"
                            for row in rows:
                                if isinstance(row, list):
                                    doc_content += " | ".join(str(cell) for cell in row) + "\n"
                                row_count += 1
                # 如果有标题结构，也添加上下文
                if doc.structured_data.get("titles"):
                    titles = doc.structured_data.get("titles", [])
                    doc_content += f"\n【文档章节结构】\n"
                    for title in titles[:20]:  # 限制前20个标题
                        doc_content += f"{'#' * title.get('level', 1)} {title.get('text', '')}\n"
                # 如果没有提取到表格内容，使用纯文本
                if not doc_content.strip():
                    doc_content = doc.content[:5000] if doc.content else ""
            elif doc.content:
                doc_content = doc.content[:5000]

            if doc_content:
                doc_context = f"【文档: {doc.filename} ({doc.doc_type})】\n{doc_content}"
                logger.info(f"文档 {doc.filename} 上下文长度: {len(doc_context)}, 行数: {row_count}")
                if total_length + len(doc_context) <= max_length:
                    contexts.append(doc_context)
                    total_length += len(doc_context)
                else:
                    remaining = max_length - total_length
                    if remaining > 100:
                        doc_context = doc_context[:remaining] + f"\n...（内容被截断）"
                        contexts.append(doc_context)
                    logger.warning(f"上下文被截断: {doc.filename}, 总长度: {total_length + len(doc_context)}")
                    break

        result = "\n\n".join(contexts) if contexts else "（源文档内容为空）"
        logger.info(f"最终上下文长度: {len(result)}")
        return result

    async def get_template_fields_from_file(
        self,
        file_path: str,
        file_type: str = "xlsx",
        source_contents: List[dict] = None
    ) -> List[TemplateField]:
        """
        从模板文件提取字段定义

        Args:
            file_path: 模板文件路径
            file_type: 文件类型 (xlsx/xls/docx)
            source_contents: 源文档内容列表（用于 AI 生成表头）

        Returns:
            字段列表
        """
        fields = []
        if source_contents is None:
            source_contents = []

        try:
            if file_type in ["xlsx", "xls"]:
                fields = await self._get_template_fields_from_excel(file_type, file_path)
            elif file_type == "docx":
                fields = await self._get_template_fields_from_docx(file_path)

            # 检查是否需要 AI 生成表头
            # 条件：没有字段 OR 所有字段都是自动命名的（如"字段1"、"列1"、"Unnamed"开头）
            needs_ai_generation = (
                len(fields) == 0 or
                all(self._is_auto_generated_field(f.name) for f in fields)
            )

            if needs_ai_generation:
                logger.info(f"模板表头为空或自动生成，尝试 AI 生成表头... (fields={len(fields)}, source_docs={len(source_contents)})")
                ai_fields = await self._generate_fields_with_ai(file_path, file_type, source_contents)
                if ai_fields:
                    fields = ai_fields
                    logger.info(f"AI 生成表头成功: {len(fields)} 个字段")

        except Exception as e:
            logger.error(f"提取模板字段失败: {str(e)}")

        return fields

    def _is_auto_generated_field(self, name: str) -> bool:
        """检查字段名是否是自动生成的（无效表头）"""
        import re
        if not name:
            return True
        name_str = str(name).strip()
        # 匹配 "字段1", "列1", "Field1", "Column1" 等自动生成的名字
        # 或 "Unnamed: 0" 等 Excel 默认名字
        if name_str.startswith('Unnamed'):
            return True
        if re.match(r'^[列字段ColumnField]+\d+$', name_str, re.IGNORECASE):
            return True
        if name_str in ['0', '1', '2'] or name_str.startswith('0.') or name_str.startswith('1.'):
            # 纯数字或类似 "0.1" 的列名
            return True
        return False

    async def _get_template_fields_from_excel(self, file_type: str, file_path: str) -> List[TemplateField]:
        """从 Excel 模板提取字段"""
        fields = []

        try:
            import pandas as pd

            # 尝试读取 Excel 文件
            try:
                # header=0 表示第一行是表头
                df = pd.read_excel(file_path, header=0, nrows=5)
            except Exception as e:
                logger.warning(f"pandas 读取 Excel 表头失败，尝试无表头模式: {e}")
                # 如果失败，尝试不使用表头模式
                df = pd.read_excel(file_path, header=None, nrows=5)
                # 如果没有表头，使用列索引作为列名
                if df.shape[1] > 0:
                    # 检查第一行是否可以作为表头
                    first_row = df.iloc[0].tolist()
                    if all(pd.notna(v) and str(v).strip() != '' for v in first_row):
                        # 第一行有内容，作为表头
                        df.columns = [str(v) if pd.notna(v) else f"列{i}" for i, v in enumerate(first_row)]
                        df = df.iloc[1:]  # 移除表头行
                    else:
                        # 第一行不是有效表头，使用默认列名
                        df.columns = [f"列{i}" for i in range(df.shape[1])]

            logger.info(f"读取 Excel 表头: {df.shape}, 列: {list(df.columns)[:10]}")

            # 如果 DataFrame 列为空或只有默认索引，尝试其他方式
            if len(df.columns) == 0 or (len(df.columns) == 1 and df.columns[0] == 0):
                logger.warning(f"表头解析结果异常，重新解析: {df.columns}")
                # 尝试读取整个文件获取列信息
                df_full = pd.read_excel(file_path, header=None)
                if df_full.shape[1] > 0:
                    # 使用第一行作为列名
                    df = df_full
                    df.columns = [str(v) if pd.notna(v) and str(v).strip() else f"列{i}" for i, v in enumerate(df.iloc[0])]
                    df = df.iloc[1:]

            for idx, col in enumerate(df.columns):
                cell = self._column_to_cell(idx)
                col_str = str(col)
                if col_str == '0' or col_str.startswith('Unnamed'):
                    col_str = f"字段{idx+1}"

                fields.append(TemplateField(
                    cell=cell,
                    name=col_str,
                    field_type=self._infer_field_type_from_value(df[col].iloc[0] if len(df) > 0 else ""),
                    required=True,
                    hint=""
                ))

            logger.info(f"从 Excel 提取到 {len(fields)} 个字段")

        except Exception as e:
            logger.error(f"从Excel提取字段失败: {str(e)}", exc_info=True)

        return fields

    async def _get_template_fields_from_docx(self, file_path: str) -> List[TemplateField]:
        """从 Word 模板提取字段"""
        fields = []

        try:
            from docx import Document

            doc = Document(file_path)

            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]

                    # 假设第一列是字段名
                    if cells and cells[0]:
                        field_name = cells[0]
                        hint = cells[1] if len(cells) > 1 else ""

                        # 跳过空行或标题行
                        if field_name and field_name not in ["", "字段名", "名称", "项目"]:
                            fields.append(TemplateField(
                                cell=f"T{table_idx}R{row_idx}",
                                name=field_name,
                                field_type=self._infer_field_type_from_hint(hint),
                                required=True,
                                hint=hint
                            ))

        except Exception as e:
            logger.error(f"从Word提取字段失败: {str(e)}")

        return fields

    def _infer_field_type_from_hint(self, hint: str) -> str:
        """从提示词推断字段类型"""
        hint_lower = hint.lower()

        date_keywords = ["年", "月", "日", "日期", "时间", "出生"]
        if any(kw in hint for kw in date_keywords):
            return "date"

        number_keywords = ["数量", "金额", "人数", "面积", "增长", "比率", "%", "率", "总计", "合计"]
        if any(kw in hint_lower for kw in number_keywords):
            return "number"

        return "text"

    def _infer_field_type_from_value(self, value: Any) -> str:
        """从示例值推断字段类型"""
        if value is None or value == "":
            return "text"

        value_str = str(value)

        # 检查日期模式
        import re
        if re.search(r'\d{4}[年/-]\d{1,2}[月/-]\d{1,2}', value_str):
            return "date"

        # 检查数值
        try:
            float(value_str.replace(',', '').replace('%', ''))
            return "number"
        except ValueError:
            pass

        return "text"

    def _column_to_cell(self, col_idx: int) -> str:
        """将列索引转换为单元格列名 (0 -> A, 1 -> B, ...)"""
        result = ""
        while col_idx >= 0:
            result = chr(65 + (col_idx % 26)) + result
            col_idx = col_idx // 26 - 1
        return result

    def _extract_value_from_text(self, text: str, field_name: str) -> str:
        """
        从非 JSON 文本中提取字段值（单值版本）

        Args:
            text: 原始文本
            field_name: 字段名称

        Returns:
            提取的值
        """
        values = self._extract_values_from_text(text, field_name)
        return values[0] if values else ""

    def _extract_values_from_structured_data(self, source_docs: List[SourceDocument], field_name: str) -> List[str]:
        """
        从结构化数据（Excel rows 或 Markdown tables）中直接提取指定列的值

        适用于有 rows 结构的文档数据，无需 LLM 即可提取

        Args:
            source_docs: 源文档列表
            field_name: 字段名称

        Returns:
            值列表，如果无法提取则返回空列表
        """
        all_values = []
        logger.info(f"[_extract_values_from_structured_data] 开始提取字段: {field_name}")
        logger.info(f"  source_docs 数量: {len(source_docs)}")

        for doc_idx, doc in enumerate(source_docs):
            # 尝试从 structured_data 中提取
            structured = doc.structured_data
            logger.info(f"  文档[{doc_idx}]: {doc.filename}, structured类型: {type(structured)}, 是否为空: {not bool(structured)}")
            if structured:
                logger.info(f"    structured_data keys: {list(structured.keys())}")

            if not structured:
                continue

            # 处理多 sheet 格式: {sheets: {sheet_name: {columns, rows}}}
            if structured.get("sheets"):
                sheets = structured.get("sheets", {})
                for sheet_name, sheet_data in sheets.items():
                    if isinstance(sheet_data, dict):
                        columns = sheet_data.get("columns", [])
                        rows = sheet_data.get("rows", [])
                        values = self._extract_column_values(rows, columns, field_name)
                        if values:
                            all_values.extend(values)
                            logger.info(f"从 sheet {sheet_name} 提取到 {len(values)} 个值")
                            break  # 只用第一个匹配的 sheet
                if all_values:
                    break

            # 处理 Markdown 表格格式: {headers: [...], rows: [...], ...}
            elif structured.get("headers") and structured.get("rows"):
                headers = structured.get("headers", [])
                rows = structured.get("rows", [])
                values = self._extract_values_from_markdown_table(headers, rows, field_name)
                if values:
                    all_values.extend(values)
                    logger.info(f"从 Markdown 文档 {doc.filename} 提取到 {len(values)} 个值")
                    break

            # 处理 MongoDB 存储的 tables 格式: {tables: [{headers, rows, ...}, ...]}
            elif structured.get("tables") and isinstance(structured.get("tables"), list):
                tables = structured.get("tables", [])
                logger.info(f"  检测到 tables 格式，共 {len(tables)} 个表")
                for table_idx, table in enumerate(tables):
                    if isinstance(table, dict):
                        headers = table.get("headers", [])
                        rows = table.get("rows", [])
                        logger.info(f"  表格[{table_idx}]: headers={headers[:3]}..., rows数量={len(rows)}")
                        values = self._extract_values_from_markdown_table(headers, rows, field_name)
                        if values:
                            all_values.extend(values)
                            logger.info(f"从表格[{table_idx}] 提取到 {len(values)} 个值")
                            break
                if all_values:
                    break

            # 处理单 sheet 格式: {columns: [...], rows: [...]}
            elif structured.get("rows"):
                columns = structured.get("columns", [])
                rows = structured.get("rows", [])
                values = self._extract_column_values(rows, columns, field_name)
                if values:
                    all_values.extend(values)
                    logger.info(f"从文档 {doc.filename} 提取到 {len(values)} 个值")
                    break

            # 处理 Markdown 表格格式: {tables: [{headers: [...], rows: [...]}]}
            elif structured.get("tables"):
                tables = structured.get("tables", [])
                for table in tables:
                    if isinstance(table, dict):
                        headers = table.get("headers", [])
                        rows = table.get("rows", [])
                        values = self._extract_column_values(rows, headers, field_name)
                        if values:
                            all_values.extend(values)
                            logger.info(f"从 Markdown 表格提取到 {len(values)} 个值")
                            break
                if all_values:
                    break

            # 处理 AI 解析的 Word 文档键值对格式: {key_values: {"键": "值"}, ...}
            if structured.get("key_values") and isinstance(structured.get("key_values"), dict):
                key_values = structured.get("key_values", {})
                logger.info(f"  检测到 AI 解析键值对格式，共 {len(key_values)} 个键值对")
                values = self._extract_from_key_values(key_values, field_name)
                if values:
                    all_values.extend(values)
                    logger.info(f"从 Word AI 键值对提取到 {len(values)} 个值: {values}")
                    break

            # 处理 AI 解析的 list_items 格式
            if structured.get("list_items") and isinstance(structured.get("list_items"), list):
                list_items = structured.get("list_items", [])
                logger.info(f"  检测到 AI 解析列表格式，共 {len(list_items)} 个列表项")
                values = self._extract_from_list_items(list_items, field_name)
                if values:
                    all_values.extend(values)
                    logger.info(f"从 Word AI 列表提取到 {len(values)} 个值")
                    break

        # 如果从结构化数据中没有提取到值，且字段是通用表头，搜索文本内容
        if not all_values and field_name in self.GENERIC_HEADER_EXPANSION:
            for doc in source_docs:
                if doc.content:
                    text_values = self._search_generic_header_in_text(doc.content, field_name)
                    if text_values:
                        all_values.extend(text_values)
                        logger.info(f"从文本内容通过通用表头匹配提取到 {len(text_values)} 个值")
                        break

        return all_values

    def _extract_values_from_markdown_table(self, headers: List, rows: List, field_name: str) -> List[str]:
        """
        从 Markdown 表格中提取指定列的值

        Markdown 表格格式:
        - headers: ["col1", "col2", ...]
        - rows: [["val1", "val2", ...], ...]

        Args:
            headers: 表头列表
            rows: 数据行列表
            field_name: 要提取的字段名

        Returns:
            值列表
        """
        if not rows or not headers:
            logger.warning(f"Markdown 表格为空: headers={headers}, rows={len(rows) if rows else 0}")
            return []

        # 查找匹配的列索引 - 使用增强的匹配算法
        target_idx = self._find_best_matching_column(headers, field_name)

        # 如果没有找到列匹配，尝试在第一列中搜索字段名（适用于指标在行的文档）
        matched_row_idx = None
        if target_idx is None and rows:
            matched_row_idx = self._search_row_in_first_column(rows, field_name)
            if matched_row_idx is not None:
                logger.info(f"在第一列找到匹配: {field_name} -> 行索引 {matched_row_idx} (转置表格结构)")

        if target_idx is None and matched_row_idx is None:
            logger.warning(f"未找到匹配列: {field_name}, 可用表头: {headers}")
            return []

        # 如果在第一列找到匹配（转置表格），提取该行的其他列作为值
        if matched_row_idx is not None:
            matched_row = rows[matched_row_idx]
            if isinstance(matched_row, list):
                # 跳过第一列（指标名），提取后续列的值
                for val in matched_row[1:]:
                    values.append(self._format_value(val))
            logger.info(f"转置表格提取到 {len(values)} 个值: {values[:5]}...")
            return self._filter_valid_values(values)

        logger.info(f"列匹配成功: {field_name} -> {headers[target_idx]} (索引: {target_idx})")

        values = []
        for row in rows:
            if isinstance(row, list):
                # 跳过子表头行（主要包含年份值的行，如 "1985", "1995"）
                if self._is_year_subheader_row(row):
                    logger.info(f"跳过子表头行: {row[:5]}...")
                    continue
                # 跳过章节标题行
                if self._is_section_header_row(row):
                    logger.info(f"跳过章节标题行: {row[:5]}...")
                    continue
                if target_idx < len(row):
                    val = row[target_idx]
                else:
                    val = ""
            else:
                val = ""
            values.append(self._format_value(val))

        # 过滤掉无效值（章节标题、省略号等）
        valid_values = self._filter_valid_values(values)
        if len(valid_values) < len(values):
            logger.info(f"过滤无效值: {len(values)} -> {len(valid_values)}")

        return valid_values

    def _is_year_subheader_row(self, row: List) -> bool:
        """
        检测行是否看起来像年份子表头行

        年份子表头行通常包含 "1985", "1995", "2020" 等4位数字

        Args:
            row: 行数据

        Returns:
            是否是年份子表头行
        """
        if not row:
            return False

        import re
        year_pattern = re.compile(r'^(19|20)\d{2}$')

        # 计算看起来像年份的单元格数量
        year_like_count = 0
        for cell in row:
            cell_str = str(cell).strip()
            if year_pattern.match(cell_str):
                year_like_count += 1

        # 如果超过50%的单元格是年份格式，认为是子表头行
        if len(row) > 0 and year_like_count / len(row) > 0.5:
            return True

        return False

    def _is_section_header_row(self, row: List) -> bool:
        """
        检测行是否看起来像章节标题行

        章节标题行通常包含 "其中："、"全部工业中："、"按...计算" 等关键词

        Args:
            row: 行数据

        Returns:
            是否是章节标题行
        """
        if not row:
            return False

        import re
        # 章节标题通常包含这些模式
        section_patterns = [
            r'其中[:：]',
            r'全部\w+中[:：]',
            r'按\w+计算',
            r'小计',
            r'合计',
            r'总计',
            r'^其中$',
            r'全部$'
        ]

        for cell in row:
            cell_str = str(cell).strip()
            if not cell_str:
                continue
            for pattern in section_patterns:
                if re.search(pattern, cell_str):
                    return True

        return False

    def _is_valid_data_value(self, val: str) -> bool:
        """
        检测值是否是有效的数据值（不是章节标题、省略号等）

        Args:
            val: 值字符串

        Returns:
            是否是有效数据值
        """
        if not val or not str(val).strip():
            return False

        val_str = str(val).strip()

        # 无效模式
        invalid_patterns = [
            r'^…$',  # 省略号
            r'^[\.。]+$',  # 只有点或句号
            r'其中[:：]',  # 章节标题
            r'全部\w+中',  # 章节标题
            r'按\w+计算',  # 计算类型
            r'^(小计|合计|总计)$',  # 汇总行
            r'^其中$',
            r'^全部$'
        ]

        for pattern in invalid_patterns:
            import re
            if re.match(pattern, val_str):
                return False

        return True

    def _filter_valid_values(self, values: List[str]) -> List[str]:
        """
        过滤出有效的数据值

        Args:
            values: 值列表

        Returns:
            只包含有效值的列表
        """
        valid_values = []
        for val in values:
            if self._is_valid_data_value(val):
                valid_values.append(val)
        return valid_values

    def _extract_from_key_values(self, key_values: Dict[str, str], field_name: str) -> List[str]:
        """
        从键值对字典中提取与字段名匹配的值

        Args:
            key_values: 键值对字典，如 {"医院数量": "38710个", "床位总数": "456789张"}
            field_name: 要匹配的字段名

        Returns:
            匹配的值列表
        """
        if not key_values:
            return []

        field_lower = field_name.lower().strip()
        field_chars = set(field_lower.replace(" ", ""))
        field_keywords = set(field_lower.replace(" ", "").split())

        best_match_key = None
        best_match_score = 0

        for key, value in key_values.items():
            key_str = str(key).strip()
            key_lower = key_str.lower()
            key_chars = set(key_lower.replace(" ", ""))

            if not key_str or not value:
                continue

            # 策略1: 精确匹配（忽略大小写）
            if key_lower == field_lower:
                logger.info(f"键值对精确匹配: {field_name} -> {key_str}: {value}")
                return [str(value)]

            # 策略2: 子字符串匹配
            if field_lower in key_lower or key_lower in field_lower:
                score = max(len(field_lower), len(key_lower)) / min(len(field_lower) + 1, len(key_lower) + 1)
                if score > best_match_score:
                    best_match_score = score
                    best_match_key = (key_str, value)

            # 策略3: 关键词重叠匹配（适用于中文）
            key_keywords = set(key_lower.replace(" ", "").split())
            overlap = field_keywords & key_keywords
            if overlap and len(overlap) > 0:
                score = len(overlap) / max(len(field_keywords), len(key_keywords), 1)
                if score > best_match_score:
                    best_match_score = score
                    best_match_key = (key_str, value)

            # 策略4: 字符级包含匹配（适用于中文短字段）
            char_overlap = field_chars & key_chars
            if char_overlap:
                char_score = len(char_overlap) / max(len(field_chars), len(key_chars), 1)
                # 对于短字段（<=4字符），降低要求
                if len(field_chars) <= 4 and char_score >= 0.5:
                    if char_score > best_match_score:
                        best_match_score = char_score
                        best_match_key = (key_str, value)
                elif char_score > best_match_score and len(char_overlap) >= 2:
                    best_match_score = char_score
                    best_match_key = (key_str, value)

        # 降低阈值到 0.2，允许更多模糊匹配
        if best_match_score >= 0.2 and best_match_key:
            logger.info(f"键值对模糊匹配: {field_name} -> {best_match_key[0]}: {best_match_key[1]} (分数: {best_match_score:.2f})")
            return [str(best_match_key[1])]

        logger.warning(f"键值对未匹配到: {field_name}, 可用键: {list(key_values.keys())}")
        return []

    def _extract_from_list_items(self, list_items: List[str], field_name: str) -> List[str]:
        """
        从列表项中提取与字段名匹配的值

        Args:
            list_items: 列表项，如 ["医院数量: 38710个", "床位总数: 456789张", ...]
            field_name: 要匹配的字段名

        Returns:
            匹配的值列表
        """
        if not list_items:
            return []

        field_lower = field_name.lower().strip()
        field_keywords = set(field_lower.replace(" ", "").split())

        matched_values = []

        for item in list_items:
            item_str = str(item).strip()
            if not item_str:
                continue

            item_lower = item_str.lower()

            # 策略1: 检查列表项是否以字段名开头（格式如 "医院数量: 38710个"）
            if ':' in item_str or '：' in item_str:
                parts = item_str.replace('：', ':').split(':', 1)
                if len(parts) == 2:
                    key = parts[0].strip()
                    value = parts[1].strip()
                    key_lower = key.lower()

                    # 精确匹配键
                    if key_lower == field_lower:
                        logger.info(f"列表项键值精确匹配: {field_name} -> {value}")
                        return [value]

                    # 子字符串匹配
                    if field_lower in key_lower or key_lower in field_lower:
                        score = max(len(field_lower), len(key_lower)) / min(len(field_lower) + 1, len(key_lower) + 1)
                        if score >= 0.2:
                            logger.info(f"列表项键值模糊匹配: {field_name} -> {key}: {value} (分数: {score:.2f})")
                            matched_values.append(value)

                    # 关键词重叠
                    key_keywords = set(key_lower.replace(" ", "").split())
                    overlap = field_keywords & key_keywords
                    if overlap:
                        score = len(overlap) / max(len(field_keywords), len(key_keywords), 1)
                        if score >= 0.2:
                            matched_values.append(value)

            # 策略2: 直接匹配整个列表项
            if field_lower in item_lower or item_lower in field_lower:
                matched_values.append(item_str)
                continue

            # 策略3: 关键词重叠
            item_keywords = set(item_lower.replace(" ", "").split())
            overlap = field_keywords & item_keywords
            if overlap and len(overlap) >= 2:  # 至少2个关键词重叠
                score = len(overlap) / max(len(field_keywords), len(item_keywords), 1)
                if score >= 0.2:
                    matched_values.append(item_str)

        if matched_values:
            logger.info(f"列表项匹配到 {len(matched_values)} 个: {matched_values[:5]}")

        return matched_values

    def _find_best_matching_column(self, headers: List, field_name: str) -> Optional[int]:
        """
        查找最佳匹配的列索引

        使用多层匹配策略:
        1. 精确匹配（忽略大小写）
        2. 子字符串匹配（字段名在表头中，或表头在字段名中）
        3. 关键词重叠匹配（中文字符串分割后比对）
        4. 字符级包含匹配（适用于中文短字段）

        Args:
            headers: 表头列表
            field_name: 要匹配的字段名

        Returns:
            匹配的列索引，找不到返回 None
        """
        field_lower = field_name.lower().strip()
        # 对中文进行字符级拆分，增加匹配的灵活性
        field_chars = set(field_lower.replace(" ", ""))
        field_keywords = set(field_lower.replace(" ", "").split())

        best_match_idx = None
        best_match_score = 0

        for idx, header in enumerate(headers):
            header_str = str(header).strip()
            header_lower = header_str.lower()

            # 跳过空表头（第一列为空的情况）
            if not header_str:
                logger.info(f"跳过空表头列: 索引 {idx}")
                continue

            # 策略1: 精确匹配（忽略大小写）
            if header_lower == field_lower:
                logger.info(f"精确匹配: {field_name} -> {header_str}")
                return idx

            # 策略2: 子字符串匹配
            if field_lower in header_lower or header_lower in field_lower:
                # 计算匹配分数（较长匹配更优先）
                score = max(len(field_lower), len(header_lower)) / min(len(field_lower) + 1, len(header_lower) + 1)
                if score > best_match_score:
                    best_match_score = score
                    best_match_idx = idx
                continue

            # 策略3: 关键词重叠匹配（适用于中文）
            header_keywords = set(header_lower.replace(" ", "").split())
            overlap = field_keywords & header_keywords
            if overlap and len(overlap) > 0:
                score = len(overlap) / max(len(field_keywords), len(header_keywords), 1)
                if score > best_match_score:
                    best_match_score = score
                    best_match_idx = idx

            # 策略4: 字符级包含匹配（适用于中文短字段，如"医院"匹配"医院数量"）
            header_chars = set(header_lower.replace(" ", ""))
            char_overlap = field_chars & header_chars
            if char_overlap:
                # 计算字符重叠率，但要求至少有一定数量的重叠字符
                char_score = len(char_overlap) / max(len(field_chars), len(header_chars), 1)
                # 对于短字段（<=4字符），降低要求，只要有重叠且字符score较高即可
                if len(field_chars) <= 4 and char_score >= 0.5:
                    if char_score > best_match_score:
                        best_match_score = char_score
                        best_match_idx = idx
                elif char_score > best_match_score and len(char_overlap) >= 2:
                    # 对于较长字段，要求至少2个字符重叠
                    best_match_score = char_score
                    best_match_idx = idx

        # 降低阈值到 0.2，允许更多模糊匹配
        if best_match_score >= 0.2:
            logger.info(f"模糊匹配: {field_name} -> {headers[best_match_idx]} (分数: {best_match_score:.2f})")
            return best_match_idx

        return None

    def _search_row_in_first_column(self, rows: List, field_name: str) -> Optional[int]:
        """
        在表格第一列中搜索字段名（适用于指标在行的转置表格结构）

        对于某些中文统计文档，表格结构是转置的：
        - 第一列是指标名称（如"医院数量"）
        - 其他列是年份或数值

        Args:
            rows: 数据行列表
            field_name: 要搜索的字段名

        Returns:
            匹配的列索引（始终返回0，因为是第一列），如果没找到返回None
        """
        if not rows or not field_name:
            return None

        field_lower = field_name.lower().strip()
        field_chars = set(field_lower.replace(" ", ""))
        field_keywords = set(field_lower.replace(" ", "").split())

        for row_idx, row in enumerate(rows):
            if not isinstance(row, list) or len(row) == 0:
                continue

            first_cell = str(row[0]).strip()
            if not first_cell:
                continue

            first_cell_lower = first_cell.lower()

            # 精确匹配
            if first_cell_lower == field_lower:
                logger.info(f"第一列精确匹配字段: {field_name} -> {first_cell} (行{row_idx})")
                return 0

            # 子字符串匹配
            if field_lower in first_cell_lower or first_cell_lower in field_lower:
                score = max(len(field_lower), len(first_cell_lower)) / min(len(field_lower) + 1, len(first_cell_lower) + 1)
                if score >= 0.5:
                    logger.info(f"第一列模糊匹配字段: {field_name} -> {first_cell} (行{row_idx}, 分数:{score:.2f})")
                    return 0

            # 关键词重叠匹配
            first_keywords = set(first_cell_lower.replace(" ", "").split())
            overlap = field_keywords & first_keywords
            if overlap and len(overlap) >= 2:
                score = len(overlap) / max(len(field_keywords), len(first_keywords), 1)
                if score >= 0.3:
                    logger.info(f"第一列关键词匹配: {field_name} -> {first_cell} (行{row_idx}, 分数:{score:.2f})")
                    return 0

            # 字符级匹配（短字段）
            first_chars = set(first_cell_lower.replace(" ", ""))
            char_overlap = field_chars & first_chars
            if char_overlap and len(field_chars) <= 4:
                char_score = len(char_overlap) / max(len(field_chars), len(first_chars), 1)
                if char_score >= 0.5:
                    logger.info(f"第一列字符匹配: {field_name} -> {first_cell} (行{row_idx}, 分数:{char_score:.2f})")
                    return 0

        return None

    def _extract_column_values(self, rows: List, columns: List, field_name: str) -> List[str]:
        """
        从 rows 和 columns 中提取指定列的值

        Args:
            rows: 行数据列表
            columns: 列名列表
            field_name: 要提取的字段名

        Returns:
            值列表
        """
        if not rows or not columns:
            return []

        # 使用增强的匹配算法查找最佳匹配的列索引
        target_idx = self._find_best_matching_column(columns, field_name)

        if target_idx is None:
            logger.warning(f"未找到匹配列: {field_name}, 可用列: {columns}")
            return []

        target_col = columns[target_idx]
        logger.info(f"列匹配成功: {field_name} -> {target_col} (索引: {target_idx})")

        values = []
        for row in rows:
            # 跳过子表头行（主要包含年份值的行，如 "1985", "1995"）
            if isinstance(row, list) and self._is_year_subheader_row(row):
                continue
            # 跳过章节标题行
            if isinstance(row, list) and self._is_section_header_row(row):
                continue
            if isinstance(row, dict):
                val = row.get(target_col, "")
            elif isinstance(row, list) and target_idx < len(row):
                val = row[target_idx]
            else:
                val = ""
            values.append(self._format_value(val))

        # 过滤掉无效值（章节标题、省略号等）
        valid_values = self._filter_valid_values(values)
        if len(valid_values) < len(values):
            logger.info(f"过滤无效值: {len(values)} -> {len(valid_values)}")

        return valid_values

    def _format_value(self, val: Any) -> str:
        """
        格式化值为字符串，保持原始格式

        - 如果是浮点数但实际上等于整数，返回整数格式（如 3.0 -> "3"）
        - 如果是浮点数且有小数部分，保留小数（如 3.5 -> "3.5"）
        - 如果是整数，直接返回（如 3 -> "3"）
        - 其他类型直接转为字符串

        Args:
            val: 原始值

        Returns:
            格式化后的字符串
        """
        if val is None:
            return ""

        # 如果已经是字符串
        if isinstance(val, str):
            return val.strip()

        # 如果是布尔值
        if isinstance(val, bool):
            return "true" if val else "false"

        # 如果是数字
        if isinstance(val, (int, float)):
            # 检查是否是浮点数但等于整数
            if isinstance(val, float):
                # 检查是否是小数部分为0
                if val == int(val):
                    return str(int(val))
                else:
                    # 去除尾部多余的0，但保留必要的小数位
                    formatted = f"{val:.10f}".rstrip('0').rstrip('.')
                    return formatted
            else:
                return str(val)

        return str(val)

    def _search_generic_header_in_text(self, text: str, field_name: str) -> List[str]:
        """
        从文本中搜索通用表头对应的具体值

        例如：表头"机构" -> 搜索文本中的"医院"、"学校"、"企业"等

        Args:
            text: 文档文本内容
            field_name: 字段名称（可能是通用表头）

        Returns:
            匹配到的值列表
        """
        import re

        # 检查是否是通用表头
        generic_terms = self.GENERIC_HEADER_EXPANSION.get(field_name, [])
        if not generic_terms:
            return []

        matched_values = []

        for term in generic_terms:
            # 搜索 term + 数字/量词 的模式，如 "医院 100所"
            patterns = [
                rf'{re.escape(term)}[\s\d所个家级人万元亿元%‰]+',  # 医院100所, 企业50家
                rf'{re.escape(term)}[：:\s]+(\d+[\d。,，]?\d*)',   # 医院：100
                rf'(\d+[\d。,，]?\d*)[^\d]*{re.escape(term)}',    # 100家医院
            ]
            for pattern in patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for match in matches:
                    val = match.strip() if isinstance(match, str) else match
                    if val and len(str(val)) < 100:
                        matched_values.append(str(val))

        # 去重并保持顺序
        seen = set()
        unique_values = []
        for v in matched_values:
            if v not in seen:
                seen.add(v)
                unique_values.append(v)

        if unique_values:
            logger.info(f"通用表头 '{field_name}' 匹配到值: {unique_values[:10]}")

        return unique_values

    def _extract_values_from_json(self, result) -> List[str]:
        """
        从解析后的 JSON 对象/数组中提取值数组

        Args:
            result: json.loads() 返回的对象

        Returns:
            值列表
        """
        if isinstance(result, dict):
            # 优先找 values 数组
            if "values" in result and isinstance(result["values"], list):
                vals = [self._format_value(v).strip() for v in result["values"] if self._format_value(v).strip()]
                if vals:
                    return vals
            # 尝试找 value 字段
            if "value" in result:
                val = self._format_value(result["value"]).strip()
                if val:
                    return [val]
            # 尝试找任何数组类型的键
            for key in result.keys():
                val = result[key]
                if isinstance(val, list) and len(val) > 0:
                    if all(isinstance(v, (str, int, float, bool)) or v is None for v in val):
                        vals = [self._format_value(v).strip() for v in val if v is not None and self._format_value(v).strip()]
                        if vals:
                            return vals
                elif isinstance(val, (str, int, float, bool)):
                    return [self._format_value(val).strip()]
        elif isinstance(result, list):
            vals = [self._format_value(v).strip() for v in result if v is not None and self._format_value(v).strip()]
            if vals:
                return vals
        return []

    def _fix_json(self, json_text: str) -> str:
        """
        尝试修复损坏的 JSON 字符串

        Args:
            json_text: 原始 JSON 文本

        Returns:
            修复后的 JSON 文本，如果无法修复则返回空字符串
        """
        import re

        # 如果以 { 开头，尝试找到配对的 }
        if json_text.startswith('{'):
            # 统计括号深度
            depth = 0
            end_pos = -1
            for i, c in enumerate(json_text):
                if c == '{':
                    depth += 1
                elif c == '}':
                    depth -= 1
                    if depth == 0:
                        end_pos = i + 1
                        break

            if end_pos > 0:
                fixed = json_text[:end_pos]
                logger.info(f"修复 JSON (配对括号): {fixed[:200]}")
                return fixed

            # 如果找不到配对，尝试移除 trailing comma 和其他问题
            # 移除末尾多余的逗号
            fixed = re.sub(r',\s*([}\]])', r'\1', json_text)
            # 确保以 } 结尾
            fixed = fixed.strip()
            if fixed and not fixed.endswith('}') and not fixed.endswith(']'):
                # 尝试补全
                if fixed.startswith('{') and not fixed.endswith('}'):
                    fixed = fixed + '}'
                elif fixed.startswith('[') and not fixed.endswith(']'):
                    fixed = fixed + ']'
            logger.info(f"修复 JSON (正则): {fixed[:200]}")
            return fixed

        # 如果以 [ 开头
        elif json_text.startswith('['):
            depth = 0
            end_pos = -1
            for i, c in enumerate(json_text):
                if c == '[':
                    depth += 1
                elif c == ']':
                    depth -= 1
                    if depth == 0:
                        end_pos = i + 1
                        break

            if end_pos > 0:
                fixed = json_text[:end_pos]
                logger.info(f"修复 JSON (数组配对): {fixed[:200]}")
                return fixed

        return ""

    def _extract_values_by_regex(self, text: str) -> List[str]:
        """
        使用正则从损坏/不完整的 JSON 文本中提取 values 数组

        即使 JSON 被截断，只要能看到 "values": [...] 就能提取

        Args:
            text: 原始文本

        Returns:
            值列表
        """
        import re

        # 方法1: 查找 "values": [ 开始的位置
        values_start = re.search(r'"values"\s*:\s*\[', text)
        if values_start:
            # 从 [ 之后开始提取内容
            start_pos = values_start.end()
            remaining = text[start_pos:]

            # 提取所有被双引号包裹的字符串值
            # 使用简单正则：匹配 "..." 捕获引号内的内容
            values = re.findall(r'"([^"]+)"', remaining)

            if values:
                # 过滤掉空字符串和很短的（可能是键名）
                filtered = [v.strip() for v in values if v.strip() and len(v) > 1]
                if filtered:
                    logger.info(f"正则提取到 {len(filtered)} 个值: {filtered[:3]}")
                    return filtered

        # 方法2: 备选 - 直接查找所有 : "value" 格式的值
        all_strings = re.findall(r':\s*"([^"]{1,200})"', text)
        if all_strings:
            filtered = [s for s in all_strings if s and len(s) < 500]
            if filtered:
                logger.info(f"备选正则提取到 {len(filtered)} 个值: {filtered[:3]}")
                return filtered

        return []

    def _extract_values_from_text(self, text: str, field_name: str) -> List[str]:
        """
        从非 JSON 文本中提取多个字段值

        Args:
            text: 原始文本
            field_name: 字段名称

        Returns:
            提取的值列表
        """
        import re
        import json

        # 先尝试解析整个文本为 JSON，检查是否包含嵌套的 values 数组
        cleaned_text = text.strip()
        # 移除可能的 markdown 代码块标记
        cleaned_text = cleaned_text.replace('```json', '').replace('```', '').strip()

        try:
            # 尝试解析整个文本为 JSON
            parsed = json.loads(cleaned_text)
            if isinstance(parsed, dict):
                # 如果是 {"values": [...]} 格式，提取 values
                if "values" in parsed and isinstance(parsed["values"], list):
                    return [self._format_value(v).strip() for v in parsed["values"] if self._format_value(v).strip()]
                # 如果是其他 dict 格式，尝试找 values 键
                for key in ["values", "value", "data", "result"]:
                    if key in parsed and isinstance(parsed[key], list):
                        return [self._format_value(v).strip() for v in parsed[key] if self._format_value(v).strip()]
                    elif key in parsed:
                        return [self._format_value(parsed[key]).strip()]
            elif isinstance(parsed, list):
                return [self._format_value(v).strip() for v in parsed if self._format_value(v).strip()]
        except (json.JSONDecodeError, TypeError):
            pass

        # 尝试匹配 JSON 数组格式
        array_match = re.search(r'\[[\s\S]*?\]', text)
        if array_match:
            try:
                arr = json.loads(array_match.group())
                if isinstance(arr, list):
                    # 检查数组元素是否是 {"values": [...]} 结构
                    if arr and isinstance(arr[0], dict) and "values" in arr[0]:
                        # 提取嵌套的 values
                        result = []
                        for item in arr:
                            if isinstance(item, dict) and "values" in item and isinstance(item["values"], list):
                                result.extend([self._format_value(v).strip() for v in item["values"] if self._format_value(v).strip()])
                            elif isinstance(item, dict):
                                result.append(str(item))
                            else:
                                result.append(self._format_value(item))
                        if result:
                            return result
                    return [self._format_value(v).strip() for v in arr if self._format_value(v).strip()]
            except:
                pass

        # 尝试用分号分割（如果文本中有分号分隔的多个值）
        if '；' in text or ';' in text:
            separator = '；' if '；' in text else ';'
            parts = text.split(separator)
            values = []
            for part in parts:
                part = part.strip()
                if part and len(part) < 500:
                    # 清理 Markdown 格式
                    part = re.sub(r'^\*\*|\*\*$', '', part)
                    part = re.sub(r'^\*|\*$', '', part)
                    values.append(part.strip())
            if values:
                return values

        # 尝试多种模式匹配
        patterns = [
            # "字段名: 值" 或 "字段名：值" 格式
            rf'{re.escape(field_name)}[：:]\s*(.+?)(?:\n|$)',
            # "值" 在引号中
            rf'"value"\s*:\s*"([^"]+)"',
            # "值" 在单引号中
            rf"['\"]?value['\"]?\s*:\s*['\"]([^'\"]+)['\"]",
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                value = match.group(1).strip()
                # 清理 Markdown 格式
                value = re.sub(r'^\*\*|\*\*$', '', value)
                value = re.sub(r'^\*|\*$', '', value)
                value = value.strip()
                if value and len(value) < 1000:
                    return [value]

        # 如果无法匹配，返回原始内容
        content = text.strip()[:500] if text.strip() else ""
        return [content] if content else []

    async def _analyze_unstructured_docs_for_fields(
        self,
        source_docs: List[SourceDocument],
        field: TemplateField,
        user_hint: Optional[str] = None
    ) -> Optional[FillResult]:
        """
        对非结构化文档进行 AI 分析，尝试提取结构化数据

        适用于 Markdown 等没有表格格式的文档，通过 AI 分析提取结构化信息

        Args:
            source_docs: 源文档列表
            field: 字段定义
            user_hint: 用户提示

        Returns:
            FillResult 如果提取成功，否则返回 None
        """
        # 找出非结构化的 Markdown/TXT 文档（没有表格的）
        unstructured_docs = []
        for doc in source_docs:
            if doc.doc_type in ["md", "txt", "markdown"]:
                # 检查是否有表格
                has_tables = (
                    doc.structured_data and
                    doc.structured_data.get("tables") and
                    len(doc.structured_data.get("tables", [])) > 0
                )
                if not has_tables:
                    unstructured_docs.append(doc)

        if not unstructured_docs:
            return None

        logger.info(f"发现 {len(unstructured_docs)} 个非结构化文档，尝试 AI 分析...")

        # 对每个非结构化文档进行 AI 分析
        for doc in unstructured_docs:
            try:
                # 使用 markdown_ai_service 的 statistics 分析类型
                # 这种类型专门用于政府统计公报等包含数据的文档
                hint_text = field.hint if field.hint else f"请提取{field.name}的信息"
                if user_hint:
                    hint_text = f"{user_hint}。{hint_text}"

                # 构建查询文本
                query_text = f"{field.name} {hint_text}"

                # 使用 RAG 向量检索获取相关内容块
                rag_results = rag_service.retrieve(
                    query=query_text,
                    top_k=5,
                    min_score=0.3
                )

                # 构建上下文：优先使用 RAG 检索结果，如果检索不到则使用原始内容
                if rag_results:
                    # 使用 RAG 检索到的相关块
                    context_parts = []
                    for result in rag_results:
                        if result.get("doc_id", "").startswith(doc.doc_id) or not result.get("doc_id"):
                            context_parts.append(result["content"])

                    if context_parts:
                        retrieved_context = "\n\n---\n\n".join(context_parts)
                        logger.info(f"RAG 检索到 {len(context_parts)} 个相关块用于字段 {field.name}")
                        # 使用检索到的内容（限制长度）
                        context_to_use = retrieved_context[:6000]
                    else:
                        # RAG 检索结果不属于当前文档，使用原始内容
                        context_to_use = doc.content[:6000] if doc.content else ""
                        logger.info(f"字段 {field.name} 使用原始内容（RAG结果不属于当前文档）")
                else:
                    # 没有 RAG 检索结果，使用原始内容
                    context_to_use = doc.content[:6000] if doc.content else ""
                    logger.info(f"字段 {field.name} 使用原始内容（无RAG检索结果）")

                # 构建针对字段提取的提示词 - 增强语义匹配能力
                prompt = f"""你是一个专业的数据提取专家。请从以下文档内容中进行**语义匹配**提取。

【重要】字段名: "{field.name}"
【重要】字段提示: {hint_text}

## 分类数据识别

文档中经常包含分类统计数据，格式如下：

### 1. 直接分类（用"其中："、"中，"等分隔）
原文示例：
- "全国医疗卫生机构总数1093551个，其中：医院38710个，基层医疗卫生机构1040023个"
  → 字段"医院数量" 应提取: 38710
  → 字段"基层医疗卫生机构数量" 应提取: 1040023

- "医院中，公立医院11754个，民营医院26956个"
  → 字段"公立医院数量" 应提取: 11754
  → 字段"民营医院数量" 应提取: 26956

### 2. 嵌套分类（用"按...分："、"其中："等结构）
原文示例：
- "医院按等级分：三级医院4111个（其中：三级甲等医院1876个），二级医院12294个"
  → 字段"三级医院数量" 应提取: 4111
  → 字段"三级甲等医院数量" 应提取: 1876
  → 字段"二级医院数量" 应提取: 12294

### 3. 匹配技巧
- "医院数量" 可匹配: "医院38710个"、"医院数量为"
- "公立医院数量" 可匹配: "公立医院11754个"、"公立医院有"
- 忽略"数量"、"数"、"个"等后缀的差异
- 数值可能紧跟关键词，也可能分开描述

## 提取规则

1. **全文搜索**：在文档的全部内容中搜索，不要只搜索开头部分
2. **分类定位**：找到包含该分类关键词的句子，理解其完整的数值
3. **保留单位**：提取数值时**要包含单位**

【重要】返回值规则：
- **返回数值时必须包含单位**
- 例如原文"公共图书馆3246个"，提取时应返回 "3246个"
- 例如原文"国内旅游收入4.9万亿元"，提取时应返回 "4.9万亿元"
- 例如原文"注册护士585.5万人"，提取时应返回 "585.5万人"
- 如果字段是"指标"类型，返回具体的指标名称文本（不带单位）
- 如果没有找到任何相关数据，返回空数组

文档内容：
{context_to_use}

请用严格的 JSON 格式返回：
{{
    "values": ["提取到的值1", "值2", ...],
    "source": "数据来源说明（如：从文档第X段提取）",
    "confidence": 0.0到1.0之间的置信度
}}

【重要】即使是模糊匹配，也要：
- 确保提取的内容确实来自文档
- source中准确说明数据来源位置"""

                messages = [
                    {"role": "system", "content": "你是一个专业的数据提取助手，擅长从政府统计公报等文档中提取数据。请严格按JSON格式输出。"},
                    {"role": "user", "content": prompt}
                ]

                response = await self.llm.chat(
                    messages=messages,
                    temperature=0.1,
                    max_tokens=4000
                )

                content = self.llm.extract_message_content(response)
                logger.info(f"AI 分析返回: {content[:500]}")

                # 解析 JSON
                import json
                import re

                # 清理 markdown 格式
                cleaned = content.strip()
                cleaned = re.sub(r'^```json\s*', '', cleaned, flags=re.MULTILINE)
                cleaned = re.sub(r'^```\s*', '', cleaned, flags=re.MULTILINE)
                cleaned = cleaned.strip()

                # 查找 JSON
                json_start = -1
                for i, c in enumerate(cleaned):
                    if c == '{' or c == '[':
                        json_start = i
                        break

                if json_start == -1:
                    continue

                json_text = cleaned[json_start:]
                try:
                    result = json.loads(json_text)
                    values = self._extract_values_from_json(result)
                    if values:
                        return FillResult(
                            field=field.name,
                            values=values,
                            value=values[0] if values else "",
                            source=f"AI分析: {doc.filename}",
                            confidence=result.get("confidence", 0.8)
                        )
                except json.JSONDecodeError:
                    # 尝试修复 JSON
                    fixed = self._fix_json(json_text)
                    if fixed:
                        try:
                            result = json.loads(fixed)
                            values = self._extract_values_from_json(result)
                            if values:
                                return FillResult(
                                    field=field.name,
                                    values=values,
                                    value=values[0] if values else "",
                                    source=f"AI分析: {doc.filename}",
                                    confidence=result.get("confidence", 0.8)
                                )
                        except json.JSONDecodeError:
                            pass

            except Exception as e:
                logger.warning(f"AI 分析文档 {doc.filename} 失败: {str(e)}")
                continue

        return None

    async def _generate_fields_with_ai(
        self,
        file_path: str,
        file_type: str,
        source_contents: List[dict] = None
    ) -> Optional[List[TemplateField]]:
        """
        使用 AI 为空表生成表头字段

        当模板文件为空或没有表头时，调用 AI 分析并生成合适的字段名

        Args:
            file_path: 模板文件路径
            file_type: 文件类型

        Returns:
            生成的字段列表，如果失败返回 None
        """
        try:
            import pandas as pd

            # 读取 Excel 内容检查是否为空
            content_sample = ""
            if file_type in ["xlsx", "xls"]:
                df = pd.read_excel(file_path, header=None)
                if df.shape[0] == 0 or df.shape[1] == 0:
                    logger.info("Excel 表格为空")
                    # 即使 Excel 为空，如果有源文档，仍然尝试使用 AI 生成表头
                    if not source_contents:
                        logger.info("Excel 为空且没有源文档，使用默认字段名")
                        return [TemplateField(
                            cell=self._column_to_cell(i),
                            name=f"字段{i+1}",
                            field_type="text",
                            required=False,
                            hint="请填写此字段"
                        ) for i in range(5)]
                    # 有源文档，继续调用 AI 生成表头
                    logger.info("Excel 为空但有源文档，使用源文档内容生成表头...")
                else:
                    # 表格有数据但没有表头
                    if df.shape[1] > 0:
                        # 读取第一行作为参考，看是否为空
                        first_row = df.iloc[0].tolist() if len(df) > 0 else []
                        if not any(pd.notna(v) and str(v).strip() != '' for v in first_row):
                            # 第一行为空，AI 生成表头
                            content_sample = df.iloc[:10].to_string() if len(df) >= 10 else df.to_string()
                        else:
                            content_sample = df.to_string()
                    else:
                        content_sample = ""

            # 调用 AI 生成表头
            # 根据源文档内容生成表头
            source_info = ""
            logger.info(f"[DEBUG] _generate_fields_with_ai received source_contents: {len(source_contents) if source_contents else 0} items")
            if source_contents:
                for sc in source_contents:
                    logger.info(f"[DEBUG]   source doc: filename={sc.get('filename')}, content_len={len(sc.get('content', ''))}, titles={len(sc.get('titles', []))}, tables_count={sc.get('tables_count', 0)}, has_tables_summary={bool(sc.get('tables_summary'))}")
                source_info = "\n\n【源文档内容摘要】（根据以下文档内容生成表头）：\n"
                for idx, src in enumerate(source_contents[:5]):  # 最多5个源文档
                    filename = src.get("filename", f"文档{idx+1}")
                    doc_type = src.get("doc_type", "unknown")
                    content = src.get("content", "")[:3000]  # 限制内容长度
                    titles = src.get("titles", [])[:10]  # 最多10个标题
                    tables_count = src.get("tables_count", 0)
                    tables_summary = src.get("tables_summary", "")

                    source_info += f"\n--- 文档 {idx+1}: {filename} ({doc_type}) ---\n"
                    # 处理 titles（可能是字符串列表或字典列表）
                    if titles:
                        title_texts = []
                        for t in titles[:5]:
                            if isinstance(t, dict):
                                title_texts.append(t.get('text', ''))
                            else:
                                title_texts.append(str(t))
                        if title_texts:
                            source_info += f"【章节标题】: {', '.join(title_texts)}\n"
                    if tables_count > 0:
                        source_info += f"【包含表格数】: {tables_count}\n"
                    if tables_summary:
                        source_info += f"{tables_summary}\n"
                    if content:
                        source_info += f"【文档内容】（前3000字符）：{content[:3000]}\n"

            prompt = f"""你是一个专业的数据分析助手。请分析源文档中的所有数据，生成表格表头字段。

任务：分析源文档，找出所有具体的数据指标及其分类。

{source_info}

【重要要求】
1. **只生成数据字段名**：
   - ✅ 正确示例："医院数量"、"公立医院数量"、"病床使用率"
   - ❌ 错误示例："source"、"备注"、"说明"、"数据来源"

2. **识别所有数值数据**：
   - 例如："医院38710个"、"病床使用率78.8%"
   - 例如："公立医院11754个"、"公立医院病床使用率84.8%"

3. **理解分类层级**：
   - 顶级分类：如"医院"、"基层医疗卫生机构"
   - 二级分类：如"医院"下分为"公立医院"、"民营医院"

4. **生成字段**：
   - 字段名要详细具体，能区分不同数据，如："医院数量（个）"、"病床使用率（%）"、"公立医院数量"
   - 优先选择：总数 + 主要分类 + 重要指标

5. **生成数量**：
   - 生成10-15个最有代表性的字段，确保覆盖主要数据指标

6. **添加字段说明**：
   - 每个字段可以添加 hint 说明字段的含义和数据来源

请严格按照以下 JSON 格式输出（只需输出 JSON，不要其他内容）：
{{
    "fields": [
        {{"name": "医院数量", "hint": "从文档中提取医院总数，包括公立和民营医院"}},
        {{"name": "病床使用率", "hint": "提取病床使用率数据"}}
    ]
}}
"""
            messages = [
                {"role": "system", "content": "你是一个专业的表格设计助手。请严格按JSON格式输出，为每个字段生成详细名称和hint说明。"},
                {"role": "user", "content": prompt}
            ]

            response = await self.llm.chat(
                messages=messages,
                temperature=0.3,
                max_tokens=4000
            )

            content = self.llm.extract_message_content(response)
            logger.info(f"AI 生成表头返回: {content[:500]}")

            # 解析 JSON
            import json
            import re

            # 清理 markdown 格式
            cleaned = content.strip()
            cleaned = re.sub(r'^```json\s*', '', cleaned, flags=re.MULTILINE)
            cleaned = re.sub(r'^```\s*', '', cleaned, flags=re.MULTILINE)
            cleaned = cleaned.strip()

            # 查找 JSON
            json_start = -1
            for i, c in enumerate(cleaned):
                if c == '{':
                    json_start = i
                    break

            if json_start == -1:
                logger.warning("无法找到 JSON 开始位置")
                return None

            json_text = cleaned[json_start:]
            result = json.loads(json_text)

            if result and "fields" in result:
                fields = []
                # 过滤非数据字段
                skip_keywords = ["source", "来源", "备注", "说明", "备注列", "说明列", "data_source", "remark", "note", "description"]
                for idx, f in enumerate(result["fields"]):
                    field_name = f.get("name", f"字段{idx+1}")
                    # 跳过非数据字段
                    if any(kw in field_name.lower() for kw in skip_keywords):
                        logger.info(f"跳过非数据字段: {field_name}")
                        continue
                    fields.append(TemplateField(
                        cell=self._column_to_cell(idx),
                        name=field_name,
                        field_type="text",
                        required=False,
                        hint=f.get("hint", "")
                    ))
                logger.info(f"AI 生成表头: {[f.name for f in fields]}")
                return fields

        except Exception as e:
            logger.error(f"AI 生成表头失败: {str(e)}")

        return None


# ==================== 全局单例 ====================

template_fill_service = TemplateFillService()
