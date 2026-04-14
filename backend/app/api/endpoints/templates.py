"""
表格模板 API 接口

提供模板上传、解析和填写功能
"""
import io
import logging
import uuid
from typing import List, Optional

from fastapi import APIRouter, File, HTTPException, Query, UploadFile, BackgroundTasks
from fastapi.responses import StreamingResponse
import pandas as pd
from pydantic import BaseModel

from app.services.template_fill_service import template_fill_service, TemplateField
from app.services.file_service import file_service
from app.core.database import mongodb
from app.core.document_parser import ParserFactory

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/templates", tags=["表格模板"])


# ==================== 辅助函数 ====================

async def update_task_status(
    task_id: str,
    status: str,
    progress: int = 0,
    message: str = "",
    result: dict = None,
    error: str = None
):
    """
    更新任务状态，同时写入 Redis 和 MongoDB
    """
    from app.core.database import redis_db

    meta = {"progress": progress, "message": message}
    if result:
        meta["result"] = result
    if error:
        meta["error"] = error

    try:
        await redis_db.set_task_status(task_id, status, meta)
    except Exception as e:
        logger.warning(f"Redis 任务状态更新失败: {e}")

    try:
        await mongodb.update_task(
            task_id=task_id,
            status=status,
            message=message,
            result=result,
            error=error
        )
    except Exception as e:
        logger.warning(f"MongoDB 任务状态更新失败: {e}")


# ==================== 请求/响应模型 ====================

class TemplateFieldRequest(BaseModel):
    """模板字段请求"""
    cell: str
    name: str
    field_type: str = "text"
    required: bool = True
    hint: str = ""


class FillRequest(BaseModel):
    """填写请求"""
    template_id: str
    template_fields: List[TemplateFieldRequest]
    source_doc_ids: Optional[List[str]] = None  # MongoDB 文档 ID 列表
    source_file_paths: Optional[List[str]] = None  # 源文档文件路径列表
    user_hint: Optional[str] = None
    task_id: Optional[str] = None  # 可选的任务ID，用于任务历史跟踪


class ExportRequest(BaseModel):
    """导出请求"""
    template_id: str
    filled_data: dict
    format: str = "xlsx"  # xlsx 或 docx


class FillAndExportRequest(BaseModel):
    """填充并导出请求 - 直接填充原始模板"""
    template_path: str  # 模板文件路径
    filled_data: dict  # 填写数据，格式: {字段名: [值1, 值2, ...]} 或 {字段名: 单个值}
    format: str = "xlsx"  # xlsx 或 docx


# ==================== 接口实现 ====================

@router.post("/upload")
async def upload_template(
    file: UploadFile = File(...),
):
    """
    上传表格模板文件

    支持 Excel (.xlsx, .xls) 和 Word (.docx) 格式

    Returns:
        模板信息，包括提取的字段列表
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名为空")

    file_ext = file.filename.split('.')[-1].lower()
    if file_ext not in ['xlsx', 'xls', 'docx']:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的模板格式: {file_ext}，仅支持 xlsx/xls/docx"
        )

    try:
        # 保存文件
        content = await file.read()
        saved_path = file_service.save_uploaded_file(
            content,
            file.filename,
            subfolder="templates"
        )

        # 提取字段
        template_fields = await template_fill_service.get_template_fields_from_file(
            saved_path,
            file_ext
        )

        return {
            "success": True,
            "template_id": saved_path,
            "filename": file.filename,
            "file_type": file_ext,
            "fields": [
                {
                    "cell": f.cell,
                    "name": f.name,
                    "field_type": f.field_type,
                    "required": f.required,
                    "hint": f.hint
                }
                for f in template_fields
            ],
            "field_count": len(template_fields)
        }

    except Exception as e:
        logger.error(f"上传模板失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"上传失败: {str(e)}")


@router.post("/upload-joint")
async def upload_joint_template(
    background_tasks: BackgroundTasks,
    template_file: UploadFile = File(..., description="模板文件"),
    source_files: List[UploadFile] = File(..., description="源文档文件列表"),
):
    """
    联合上传模板和源文档，一键完成解析和存储

    1. 保存模板文件并提取字段
    2. 异步处理源文档（解析+存MongoDB）
    3. 返回模板信息和源文档ID列表

    Args:
        template_file: 模板文件 (xlsx/xls/docx)
        source_files: 源文档列表 (docx/xlsx/md/txt)

    Returns:
        模板ID、字段列表、源文档ID列表
    """
    if not template_file.filename:
        raise HTTPException(status_code=400, detail="模板文件名为空")

    # 验证模板格式
    template_ext = template_file.filename.split('.')[-1].lower()
    if template_ext not in ['xlsx', 'xls', 'docx']:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的模板格式: {template_ext}，仅支持 xlsx/xls/docx"
        )

    # 验证源文档格式
    valid_exts = ['docx', 'xlsx', 'xls', 'md', 'txt']
    for sf in source_files:
        if sf.filename:
            sf_ext = sf.filename.split('.')[-1].lower()
            if sf_ext not in valid_exts:
                raise HTTPException(
                    status_code=400,
                    detail=f"不支持的源文档格式: {sf_ext}，仅支持 docx/xlsx/xls/md/txt"
                )

    try:
        # 1. 保存模板文件
        template_content = await template_file.read()
        template_path = file_service.save_uploaded_file(
            template_content,
            template_file.filename,
            subfolder="templates"
        )

        # 2. 保存并解析源文档 - 提取内容用于生成表头
        source_file_info = []
        source_contents = []
        for sf in source_files:
            if sf.filename:
                sf_content = await sf.read()
                sf_ext = sf.filename.split('.')[-1].lower()
                sf_path = file_service.save_uploaded_file(
                    sf_content,
                    sf.filename,
                    subfolder=sf_ext
                )
                source_file_info.append({
                    "path": sf_path,
                    "filename": sf.filename,
                    "ext": sf_ext
                })
                # 解析源文档获取内容（用于 AI 生成表头）
                try:
                    from app.core.document_parser import ParserFactory
                    parser = ParserFactory.get_parser(sf_path)
                    parse_result = parser.parse(sf_path)
                    if parse_result.success and parse_result.data:
                        # 获取原始内容
                        content = parse_result.data.get("content", "")[:5000] if parse_result.data.get("content") else ""

                        # 获取标题（可能在顶层或structured_data内）
                        titles = parse_result.data.get("titles", [])
                        if not titles and parse_result.data.get("structured_data"):
                            titles = parse_result.data.get("structured_data", {}).get("titles", [])
                        titles = titles[:10] if titles else []

                        # 获取表格数量（可能在顶层或structured_data内）
                        tables = parse_result.data.get("tables", [])
                        if not tables and parse_result.data.get("structured_data"):
                            tables = parse_result.data.get("structured_data", {}).get("tables", [])
                        tables_count = len(tables) if tables else 0

                        # 获取表格内容摘要（用于 AI 理解源文档结构）
                        tables_summary = ""
                        if tables:
                            tables_summary = "\n【文档中的表格】:\n"
                            for idx, table in enumerate(tables[:5]):  # 最多5个表格
                                if isinstance(table, dict):
                                    headers = table.get("headers", [])
                                    rows = table.get("rows", [])
                                    if headers:
                                        tables_summary += f"表格{idx+1}表头: {', '.join(str(h) for h in headers)}\n"
                                    if rows:
                                        tables_summary += f"表格{idx+1}前3行: "
                                        for row_idx, row in enumerate(rows[:3]):
                                            if isinstance(row, list):
                                                tables_summary += " | ".join(str(c) for c in row) + "; "
                                            elif isinstance(row, dict):
                                                tables_summary += " | ".join(str(row.get(h, "")) for h in headers if headers) + "; "
                                        tables_summary += "\n"

                        source_contents.append({
                            "filename": sf.filename,
                            "doc_type": sf_ext,
                            "content": content,
                            "titles": titles,
                            "tables_count": tables_count,
                            "tables_summary": tables_summary
                        })
                        logger.info(f"[DEBUG] source_contents built: filename={sf.filename}, content_len={len(content)}, titles_count={len(titles)}, tables_count={tables_count}")
                        if tables_summary:
                            logger.info(f"[DEBUG] tables_summary preview: {tables_summary[:300]}")
                except Exception as e:
                    logger.warning(f"解析源文档失败 {sf.filename}: {e}")

        # 3. 根据源文档内容生成表头
        template_fields = await template_fill_service.get_template_fields_from_file(
            template_path,
            template_ext,
            source_contents=source_contents  # 传递源文档内容
        )

        # 3. 异步处理源文档到MongoDB
        task_id = str(uuid.uuid4())
        if source_file_info:
            # 保存任务记录到 MongoDB
            try:
                await mongodb.insert_task(
                    task_id=task_id,
                    task_type="source_process",
                    status="pending",
                    message=f"开始处理 {len(source_file_info)} 个源文档"
                )
            except Exception as mongo_err:
                logger.warning(f"MongoDB 保存任务记录失败: {mongo_err}")

            background_tasks.add_task(
                process_source_documents,
                task_id=task_id,
                files=source_file_info
            )

        logger.info(f"联合上传完成: 模板={template_file.filename}, 源文档={len(source_file_info)}个")

        return {
            "success": True,
            "template_id": template_path,
            "filename": template_file.filename,
            "file_type": template_ext,
            "fields": [
                {
                    "cell": f.cell,
                    "name": f.name,
                    "field_type": f.field_type,
                    "required": f.required,
                    "hint": f.hint
                }
                for f in template_fields
            ],
            "field_count": len(template_fields),
            "source_file_paths": [f["path"] for f in source_file_info],
            "source_filenames": [f["filename"] for f in source_file_info],
            "task_id": task_id
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"联合上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"联合上传失败: {str(e)}")


async def process_source_documents(task_id: str, files: List[dict]):
    """异步处理源文档，存入MongoDB"""
    try:
        await update_task_status(
            task_id, status="processing",
            progress=0, message="开始处理源文档"
        )

        doc_ids = []
        for i, file_info in enumerate(files):
            try:
                parser = ParserFactory.get_parser(file_info["path"])
                result = parser.parse(file_info["path"])

                if result.success:
                    doc_id = await mongodb.insert_document(
                        doc_type=file_info["ext"],
                        content=result.data.get("content", ""),
                        metadata={
                            **result.metadata,
                            "original_filename": file_info["filename"],
                            "file_path": file_info["path"]
                        },
                        structured_data=result.data.get("structured_data")
                    )
                    doc_ids.append(doc_id)
                    logger.info(f"源文档处理成功: {file_info['filename']}, doc_id: {doc_id}")
                else:
                    logger.error(f"源文档解析失败: {file_info['filename']}, error: {result.error}")

            except Exception as e:
                logger.error(f"源文档处理异常: {file_info['filename']}, error: {str(e)}")

            progress = int((i + 1) / len(files) * 100)
            await update_task_status(
                task_id, status="processing",
                progress=progress, message=f"已处理 {i+1}/{len(files)}"
            )

        await update_task_status(
            task_id, status="success",
            progress=100, message="源文档处理完成",
            result={"doc_ids": doc_ids}
        )
        logger.info(f"所有源文档处理完成: {len(doc_ids)}个")

    except Exception as e:
        logger.error(f"源文档批量处理失败: {str(e)}")
        await update_task_status(
            task_id, status="failure",
            progress=0, message="源文档处理失败",
            error=str(e)
        )


@router.post("/fields")
async def extract_template_fields(
    template_id: str = Query(..., description="模板ID/文件路径"),
    file_type: str = Query("xlsx", description="文件类型")
):
    """
    从已上传的模板提取字段定义

    Args:
        template_id: 模板ID
        file_type: 文件类型

    Returns:
        字段列表
    """
    try:
        fields = await template_fill_service.get_template_fields_from_file(
            template_id,
            file_type
        )

        return {
            "success": True,
            "fields": [
                {
                    "cell": f.cell,
                    "name": f.name,
                    "field_type": f.field_type,
                    "required": f.required,
                    "hint": f.hint
                }
                for f in fields
            ]
        }

    except Exception as e:
        logger.error(f"提取字段失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"提取失败: {str(e)}")


@router.post("/fill")
async def fill_template(
    request: FillRequest,
):
    """
    执行表格填写

    根据提供的字段定义，从源文档中检索信息并填写

    Args:
        request: 填写请求

    Returns:
        填写结果
    """
    # 生成或使用传入的 task_id
    task_id = request.task_id or str(uuid.uuid4())

    try:
        # 创建任务记录到 MongoDB
        try:
            await mongodb.insert_task(
                task_id=task_id,
                task_type="template_fill",
                status="processing",
                message=f"开始填表任务: {len(request.template_fields)} 个字段"
            )
        except Exception as mongo_err:
            logger.warning(f"MongoDB 创建任务记录失败: {mongo_err}")

        # 更新进度 - 开始
        await update_task_status(
            task_id, "processing",
            progress=0, message="开始处理..."
        )

        # 转换字段
        fields = [
            TemplateField(
                cell=f.cell,
                name=f.name,
                field_type=f.field_type,
                required=f.required,
                hint=f.hint
            )
            for f in request.template_fields
        ]

        # 从 template_id 提取文件类型
        template_file_type = "xlsx"  # 默认类型
        if request.template_id:
            ext = request.template_id.split('.')[-1].lower()
            if ext in ["xlsx", "xls"]:
                template_file_type = "xlsx"
            elif ext == "docx":
                template_file_type = "docx"

        # 更新进度 - 准备开始填写
        await update_task_status(
            task_id, "processing",
            progress=10, message=f"准备填写 {len(fields)} 个字段..."
        )

        # 执行填写
        result = await template_fill_service.fill_template(
            template_fields=fields,
            source_doc_ids=request.source_doc_ids,
            source_file_paths=request.source_file_paths,
            user_hint=request.user_hint,
            template_id=request.template_id,
            template_file_type=template_file_type,
            task_id=task_id
        )

        # 更新为成功
        await update_task_status(
            task_id, "success",
            progress=100, message="填表完成",
            result={
                "field_count": len(fields),
                "max_rows": result.get("max_rows", 0)
            }
        )

        return {**result, "task_id": task_id}

    except Exception as e:
        # 更新为失败
        await update_task_status(
            task_id, "failure",
            progress=0, message="填表失败",
            error=str(e)
        )
        logger.error(f"填写表格失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"填写失败: {str(e)}")


@router.post("/export")
async def export_filled_template(
    request: ExportRequest,
):
    """
    导出填写后的表格

    支持 Excel (.xlsx) 和 Word (.docx) 格式

    Args:
        request: 导出请求

    Returns:
        文件流
    """
    try:
        if request.format == "xlsx":
            return await _export_to_excel(request.filled_data, request.template_id)
        elif request.format == "docx":
            return await _export_to_word(request.filled_data, request.template_id)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的导出格式: {request.format}，仅支持 xlsx/docx"
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"导出失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


async def _export_to_excel(filled_data: dict, template_id: str) -> StreamingResponse:
    """导出为 Excel 格式（支持多行）"""
    import logging
    logger = logging.getLogger(__name__)

    logger.info(f"导出填表数据: {len(filled_data)} 个字段")

    # 计算最大行数
    max_rows = 1
    for k, v in filled_data.items():
        if isinstance(v, list) and len(v) > max_rows:
            max_rows = len(v)
        logger.info(f"  {k}: {type(v).__name__} = {str(v)[:80]}")

    logger.info(f"最大行数: {max_rows}")

    # 构建多行数据
    rows_data = []
    for row_idx in range(max_rows):
        row = {}
        for col_name, values in filled_data.items():
            if isinstance(values, list):
                # 取对应行的值，不足则填空
                row[col_name] = values[row_idx] if row_idx < len(values) else ""
            else:
                # 非列表，整个值填入第一行
                row[col_name] = values if row_idx == 0 else ""
        rows_data.append(row)

    df = pd.DataFrame(rows_data)

    # 确保列顺序
    if not df.empty:
        df = df[list(filled_data.keys())]

    logger.info(f"DataFrame 形状: {df.shape}")
    logger.info(f"DataFrame 列: {list(df.columns)}")

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='填写结果')

    output.seek(0)

    filename = f"filled_template.xlsx"

    return StreamingResponse(
        io.BytesIO(output.getvalue()),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


async def _export_to_word(filled_data: dict, template_id: str) -> StreamingResponse:
    """导出为 Word 格式"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 添加标题
    title = doc.add_heading('填写结果', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加填写时间和模板信息
    from datetime import datetime
    info_para = doc.add_paragraph()
    info_para.add_run(f"模板ID: {template_id}\n").bold = True
    info_para.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_paragraph()  # 空行

    # 添加字段表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '字段名'
    header_cells[1].text = '填写值'
    header_cells[2].text = '状态'

    for field_name, field_value in filled_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = field_name
        row_cells[1].text = str(field_value) if field_value else ''
        row_cells[2].text = '已填写' if field_value else '为空'

    # 保存到 BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"filled_template.docx"

    return StreamingResponse(
        io.BytesIO(output.getvalue()),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.post("/export/excel")
async def export_to_excel(
    filled_data: dict,
    template_id: str = Query(..., description="模板ID")
):
    """
    专门导出为 Excel 格式

    Args:
        filled_data: 填写数据
        template_id: 模板ID

    Returns:
        Excel 文件流
    """
    return await _export_to_excel(filled_data, template_id)


@router.post("/export/word")
async def export_to_word(
    filled_data: dict,
    template_id: str = Query(..., description="模板ID")
):
    """
    专门导出为 Word 格式

    Args:
        filled_data: 填写数据
        template_id: 模板ID

    Returns:
        Word 文件流
    """
    return await _export_to_word(filled_data, template_id)


@router.post("/fill-and-export")
async def fill_and_export_template(
    request: FillAndExportRequest,
):
    """
    填充原始模板并导出

    直接打开原始模板文件，将数据填入模板的表格中，然后导出

    Args:
        request: 填充并导出请求

    Returns:
        填充后的模板文件流
    """
    import os

    logger.info(f"=== fill-and-export 请求 ===")
    logger.info(f"template_path: {request.template_path}")
    logger.info(f"format: {request.format}")
    logger.info(f"filled_data: {request.filled_data}")
    logger.info(f"filled_data 类型: {type(request.filled_data)}")
    logger.info(f"filled_data 键数量: {len(request.filled_data) if request.filled_data else 0}")
    logger.info(f"=========================")

    template_path = request.template_path

    # 检查模板文件是否存在
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"模板文件不存在: {template_path}")

    file_ext = os.path.splitext(template_path)[1].lower()

    try:
        if file_ext in ['.xlsx', '.xls']:
            return await _fill_and_export_excel(template_path, request.filled_data)
        elif file_ext == '.docx':
            return await _fill_and_export_word(template_path, request.filled_data)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的模板格式: {file_ext}，仅支持 xlsx/xls/docx"
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"填充模板失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"填充模板失败: {str(e)}")


async def _fill_and_export_word(template_path: str, filled_data: dict) -> StreamingResponse:
    """
    填充原始 Word 模板

    打开原始 Word 模板，找到表格，将数据填入对应单元格

    Args:
        template_path: 模板文件路径
        filled_data: 填写数据 {字段名: [值1, 值2, ...]}

    Returns:
        填充后的 Word 文件流
    """
    from docx import Document

    logger.info(f"填充 Word 模板: {template_path}")
    logger.info(f"填写数据字段: {list(filled_data.keys())}")

    # 打开原始模板
    doc = Document(template_path)

    # 找到第一个表格（比赛模板通常是第一个表格）
    if not doc.tables:
        logger.warning("Word 模板中没有表格，创建新表格")
        # 如果没有表格，创建一个
        table = doc.add_table(rows=len(filled_data) + 1, cols=2)

        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '字段名'
        header_cells[1].text = '填写值'

        # 数据行
        for idx, (field_name, values) in enumerate(filled_data.items()):
            row_cells = table.rows[idx + 1].cells
            row_cells[0].text = field_name
            if isinstance(values, list):
                row_cells[1].text = '; '.join(str(v) for v in values if v)
            else:
                row_cells[1].text = str(values) if values else ''
    else:
        # 填充第一个表格
        table = doc.tables[0]
        logger.info(f"找到表格，行数: {len(table.rows)}, 列数: {len(table.columns)}")

        # 打印表格内容（调试用）
        logger.info("=== 表格内容预览 ===")
        for row_idx, row in enumerate(table.rows[:5]):  # 只打印前5行
            row_texts = [cell.text.strip() for cell in row.cells]
            logger.info(f"  行{row_idx}: {row_texts}")
        logger.info("========================")

        # 构建字段名到列索引的映射
        field_to_col = {}
        if table.rows:
            # 假设第一行是表头
            header_row = table.rows[0]
            for col_idx, cell in enumerate(header_row.cells):
                field_name = cell.text.strip()
                if field_name:
                    field_to_col[field_name] = col_idx
                    field_to_col[field_name.lower()] = col_idx  # 忽略大小写

        logger.info(f"表头字段映射: {field_to_col}")
        logger.info(f"待填充数据字段: {list(filled_data.keys())}")

        # 填充数据
        filled_count = 0
        for field_name, values in filled_data.items():
            # 查找匹配的列
            col_idx = field_to_col.get(field_name)

            if col_idx is None:
                # 尝试模糊匹配
                for c_idx in range(len(table.columns)):
                    header_text = table.rows[0].cells[c_idx].text.strip().lower()
                    if field_name.lower() in header_text or header_text in field_name.lower():
                        col_idx = c_idx
                        logger.info(f"模糊匹配成功: '{field_name}' -> 列 {col_idx}")
                        break
                else:
                    col_idx = None

            if col_idx is not None and col_idx < len(table.columns):
                # 填充该列的所有数据行
                if isinstance(values, list):
                    value_str = '; '.join(str(v) for v in values if v)
                else:
                    value_str = str(values) if values else ''

                # 填充每一行（从第二行开始，跳过表头）
                for row_idx in range(1, min(len(table.rows), len(values) + 1) if isinstance(values, list) else 2):
                    try:
                        cell = table.rows[row_idx].cells[col_idx]
                        if isinstance(values, list) and row_idx - 1 < len(values):
                            cell.text = str(values[row_idx - 1]) if values[row_idx - 1] else ''
                        elif not isinstance(values, list):
                            if row_idx == 1:
                                cell.text = str(values) if values else ''
                    except Exception as e:
                        logger.warning(f"填充单元格失败 [{row_idx}][{col_idx}]: {e}")

                filled_count += 1
                logger.info(f"✅ 字段 '{field_name}' -> 列 {col_idx}, 值: {value_str[:50]}")
            else:
                logger.warning(f"❌ 未找到字段 '{field_name}' 对应的列")

        logger.info(f"填充完成: {filled_count}/{len(filled_data)} 个字段")

    # 保存到 BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"filled_template.docx"

    logger.info(f"Word 模板填充完成")

    return StreamingResponse(
        io.BytesIO(output.getvalue()),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


async def _fill_and_export_excel(template_path: str, filled_data: dict) -> StreamingResponse:
    """
    填充原始 Excel 模板

    打开原始 Excel 模板，找到对应列，将数据填入

    Args:
        template_path: 模板文件路径
        filled_data: 填写数据 {字段名: [值1, 值2, ...]}

    Returns:
        填充后的 Excel 文件流
    """
    from openpyxl import load_workbook
    import os

    logger.info(f"填充 Excel 模板: {template_path}")
    logger.info(f"填写数据: {list(filled_data.keys())}")

    # 检查文件是否存在
    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"模板文件不存在: {template_path}")

    # 打开原始模板
    wb = load_workbook(template_path)
    ws = wb.active  # 获取当前工作表

    logger.info(f"工作表: {ws.title}, 行数: {ws.max_row}, 列数: {ws.max_column}")

    # 读取表头行（假设第一行是表头）
    header_row = 1
    field_to_col = {}

    for col_idx in range(1, ws.max_column + 1):
        cell_value = ws.cell(row=header_row, column=col_idx).value
        if cell_value:
            field_name = str(cell_value).strip()
            field_to_col[field_name] = col_idx
            field_to_col[field_name.lower()] = col_idx  # 忽略大小写

    logger.info(f"表头字段映射: {field_to_col}")

    # 计算最大行数
    max_rows = 1
    for values in filled_data.values():
        if isinstance(values, list):
            max_rows = max(max_rows, len(values))

    # 填充数据
    for field_name, values in filled_data.items():
        # 查找匹配的列
        col_idx = field_to_col.get(field_name)

        if col_idx is None:
            # 尝试模糊匹配
            for col_idx in range(1, ws.max_column + 1):
                header_text = str(ws.cell(row=header_row, column=col_idx).value or '').strip().lower()
                if field_name.lower() in header_text or header_text in field_name.lower():
                    break
            else:
                col_idx = None

        if col_idx is not None:
            # 填充数据（从第二行开始）
            if isinstance(values, list):
                for row_idx, value in enumerate(values, start=2):
                    ws.cell(row=row_idx, column=col_idx, value=value if value else '')
            else:
                ws.cell(row=2, column=col_idx, value=values if values else '')

            logger.info(f"字段 {field_name} -> 列 {col_idx}, 值数量: {len(values) if isinstance(values, list) else 1}")
        else:
            logger.warning(f"未找到字段 {field_name} 对应的列")

    # 如果需要扩展行数
    current_max_row = ws.max_row
    if max_rows > current_max_row - 1:  # -1 是因为表头占一行
        # 扩展样式（简单复制最后一行）
        for row_idx in range(current_max_row + 1, max_rows + 2):
            for col_idx in range(1, ws.max_column + 1):
                source_cell = ws.cell(row=current_max_row, column=col_idx)
                target_cell = ws.cell(row=row_idx, column=col_idx)
                # 复制值（如果有对应数据）
                if isinstance(filled_data.get(str(ws.cell(row=1, column=col_idx).value), []), list):
                    data_idx = row_idx - 2
                    data_list = filled_data.get(str(ws.cell(row=1, column=col_idx).value), [])
                    if data_idx < len(data_list):
                        target_cell.value = data_list[data_idx]

    # 保存到 BytesIO
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    # 关闭工作簿
    wb.close()

    filename = f"filled_template.xlsx"

    logger.info(f"Excel 模板填充完成")

    return StreamingResponse(
        io.BytesIO(output.getvalue()),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ==================== Word 文档结构化字段提取接口 ====================

@router.post("/parse-word-structure")
async def parse_word_structure(
    file: UploadFile = File(...),
):
    """
    上传 Word 文档，提取结构化字段并存入数据库

    专门用于比赛场景：从 Word 表格模板中提取字段定义
    （字段名、提示词、字段类型等）并存入 MongoDB

    Args:
        file: 上传的 Word 文件

    Returns:
        提取的结构化字段信息
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名为空")

    file_ext = file.filename.split('.')[-1].lower()
    if file_ext != 'docx':
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: {file_ext}，仅支持 docx"
        )

    try:
        # 1. 保存文件
        content = await file.read()
        saved_path = file_service.save_uploaded_file(
            content,
            file.filename,
            subfolder="word_templates"
        )
        logger.info(f"Word 文件已保存: {saved_path}")

        # 2. 解析文档，提取结构化数据
        parser = ParserFactory.get_parser(saved_path)
        parse_result = parser.parse(saved_path)

        if not parse_result.success:
            raise HTTPException(status_code=400, detail=parse_result.error)

        # 3. 提取表格模板字段
        from app.core.document_parser.docx_parser import DocxParser
        docx_parser = DocxParser()
        template_fields = docx_parser.extract_template_fields_from_docx(saved_path)

        logger.info(f"从 Word 文档提取到 {len(template_fields)} 个字段")

        # 4. 提取完整的结构化信息
        template_info = docx_parser.parse_tables_for_template(saved_path)

        # 5. 存储到 MongoDB
        doc_id = await mongodb.insert_document(
            doc_type="docx",
            content=parse_result.data.get("content", ""),
            metadata={
                **parse_result.metadata,
                "original_filename": file.filename,
                "file_path": saved_path,
                "template_fields": template_fields,
                "table_count": len(template_info.get("tables", [])),
                "field_count": len(template_fields)
            },
            structured_data={
                **parse_result.data.get("structured_data", {}),
                "template_fields": template_fields,
                "template_info": template_info
            }
        )

        logger.info(f"Word 文档结构化信息已存入 MongoDB, doc_id: {doc_id}")

        # 6. 返回结果
        return {
            "success": True,
            "doc_id": doc_id,
            "filename": file.filename,
            "file_path": saved_path,
            "field_count": len(template_fields),
            "fields": template_fields,
            "tables": template_info.get("tables", []),
            "metadata": {
                "paragraph_count": parse_result.metadata.get("paragraph_count", 0),
                "table_count": parse_result.metadata.get("table_count", 0),
                "word_count": parse_result.metadata.get("word_count", 0),
                "has_tables": parse_result.metadata.get("has_tables", False)
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"解析 Word 文档结构失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"解析失败: {str(e)}")


@router.get("/word-fields/{doc_id}")
async def get_word_template_fields(
    doc_id: str,
):
    """
    根据 doc_id 获取 Word 文档的模板字段信息

    Args:
        doc_id: MongoDB 文档 ID

    Returns:
        模板字段信息
    """
    try:
        doc = await mongodb.get_document(doc_id)

        if not doc:
            raise HTTPException(status_code=404, detail=f"文档不存在: {doc_id}")

        # 从 structured_data 中提取模板字段信息
        structured_data = doc.get("structured_data", {})
        template_fields = structured_data.get("template_fields", [])
        template_info = structured_data.get("template_info", {})

        return {
            "success": True,
            "doc_id": doc_id,
            "filename": doc.get("metadata", {}).get("original_filename", ""),
            "fields": template_fields,
            "tables": template_info.get("tables", []),
            "field_count": len(template_fields),
            "metadata": doc.get("metadata", {})
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取 Word 模板字段失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取失败: {str(e)}")
